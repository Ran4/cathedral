# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx", "matplotlib", "pillow"]
# ///
"""Reframe the Nine Before Rain six-pager as a quest overview.

The source document remains in docs/codex_gdd.  This script opens that DOCX with
python-docx, scopes its copy to one quest, replaces the five embedded diagrams
whose labels implied a whole-game GDD, and exports DOCX + PDF beside this file.

Run:

    uv run --cache-dir /tmp/uv-cache \
      features/quest_secure_votes_for_a_drainage_funding_plan_before_the_rain/generate_quest_overview.py
"""
from __future__ import annotations

import os
import subprocess
import tempfile
from pathlib import Path

os.environ.setdefault("MPLBACKEND", "Agg")
os.environ.setdefault("MPLCONFIGDIR", "/tmp/nine_before_rain_quest_mpl")

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import Circle, FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


HERE = Path(__file__).resolve().parent
REPO = HERE.parent.parent
SOURCE = REPO / "docs" / "codex_gdd" / "nine_before_rain_gdd.docx"
FIG = HERE / "quest_overview_figures"
FIG.mkdir(exist_ok=True)
OUT_DOCX = HERE / "nine_before_rain_quest_overview.docx"
OUT_PDF = HERE / "nine_before_rain_quest_overview.pdf"

INK = "25211D"
PARCH = "F4EBD7"
PAPER = "FBF8F0"
BLUE = "425E6D"
PALE_BLUE = "DFE8E9"
RED = "8E3328"
PALE_RED = "F0DFD9"
OCHRE = "B58A4A"
PALE_OCHRE = "EEE2CA"
MUTED = "71685D"
LINE = "C9BCA5"
GREEN = "667B69"
BODY_FONT = "EB Garamond"
LABEL_FONT = "Lato"

_fonts = {font.name for font in font_manager.fontManager.ttflist}
FIG_SERIF = BODY_FONT if BODY_FONT in _fonts else "DejaVu Serif"
FIG_SANS = LABEL_FONT if LABEL_FONT in _fonts else "DejaVu Sans"
plt.rcParams.update(
    {
        "font.family": FIG_SERIF,
        "text.color": f"#{INK}",
        "figure.facecolor": f"#{PAPER}",
        "axes.facecolor": f"#{PAPER}",
        "savefig.facecolor": f"#{PAPER}",
    }
)


def new_figure(width: float, height: float):
    fig = plt.figure(figsize=(width, height))
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, width)
    ax.set_ylim(0, height)
    ax.axis("off")
    return fig, ax


def save_figure(fig, name: str) -> Path:
    path = FIG / f"{name}.png"
    fig.savefig(path, dpi=220, bbox_inches="tight", pad_inches=0.03)
    plt.close(fig)
    return path


def box(ax, x, y, w, h, title, body="", *, face=PARCH, edge=INK, title_color=INK,
        body_color=MUTED, title_size=9.5, body_size=7.4):
    patch = FancyBboxPatch(
        (x - w / 2, y - h / 2),
        w,
        h,
        boxstyle="round,pad=0.03,rounding_size=0.08",
        fc=f"#{face}",
        ec=f"#{edge}",
        lw=1.1,
    )
    ax.add_patch(patch)
    if body:
        ax.text(x, y + 0.13, title, ha="center", va="center", fontsize=title_size,
                fontfamily=FIG_SANS, fontweight="bold", color=f"#{title_color}")
        ax.text(x, y - 0.16, body, ha="center", va="center", fontsize=body_size,
                color=f"#{body_color}", linespacing=1.1)
    else:
        ax.text(x, y, title, ha="center", va="center", fontsize=title_size,
                fontfamily=FIG_SANS, fontweight="bold", color=f"#{title_color}")
    return patch


def arrow(ax, start, end, *, color=BLUE, width=1.25):
    patch = FancyArrowPatch(
        start,
        end,
        arrowstyle="-|>",
        mutation_scale=10,
        linewidth=width,
        color=f"#{color}",
        shrinkA=7,
        shrinkB=7,
    )
    ax.add_patch(patch)


def quest_loop_figure() -> Path:
    fig, ax = new_figure(7.2, 3.45)
    nodes = [
        (1.15, 2.55, "READ", "forecast, bells,\nquest receipts"),
        (3.60, 2.88, "CHOOSE", "one or two leads;\nchange the draft"),
        (6.05, 2.55, "FIND", "learn a schedule;\nwalk there in time"),
        (6.05, 0.82, "DO", "carry, inspect, trade,\nmark, trespass"),
        (3.60, 0.48, "RETURN", "hear the condition;\nwin a named seal"),
        (1.15, 0.82, "SHELTER / END DAY", "sleep, or spend the\nnight breaking rules"),
    ]
    for x, y, title, body in nodes:
        box(ax, x, y, 1.8 if title != "SHELTER / END DAY" else 2.1, 0.72, title, body,
            face=PAPER, edge=BLUE, title_color=BLUE, title_size=8.8 if title == "SHELTER / END DAY" else 9.5)
    connectors = [
        ((2.05, 2.67), (2.70, 2.76)),
        ((4.50, 2.76), (5.15, 2.67)),
        ((6.05, 2.19), (6.05, 1.18)),
        ((5.15, 0.70), (4.50, 0.59)),
        ((2.70, 0.59), (2.20, 0.70)),
        ((1.15, 1.18), (1.15, 2.19)),
    ]
    for start, end in connectors:
        arrow(ax, start, end, width=1.45)
    ax.add_patch(Circle((3.60, 1.66), 0.62, fc=f"#{PALE_RED}", ec=f"#{RED}", lw=1.4))
    ax.text(3.60, 1.81, "9 OF 16", ha="center", va="center", fontsize=14,
            fontfamily=FIG_SANS, fontweight="bold", color=f"#{RED}")
    ax.text(3.60, 1.50, "before hard rain", ha="center", va="center", fontsize=8.5,
            fontstyle="italic", color=f"#{RED}")
    ax.text(0.08, 1.68, "one quest day", rotation=90, ha="center", va="center",
            fontsize=7.5, fontfamily=FIG_SANS, fontweight="bold", color=f"#{MUTED}")
    return save_figure(fig, "quest_loop")


def quest_arc_figure() -> Path:
    fig, ax = new_figure(7.2, 2.15)
    days = [
        ("DAY 1", "THE WRIT", "learn one branch\nfirst named seal", "1–2"),
        ("DAY 2", "THE PRICE", "work becomes terms\npromises conflict", "3–4"),
        ("DAY 3", "THE TURN", "rumor returns\nfirst withdrawal", "5–6"),
        ("DAY 4", "THE SHORTFALL", "riders, leverage,\ndirty shortcuts", "7–8"),
        ("DAY 5", "THE HAND-COUNT", "final words at High Wick\nrain after Lamplight", "9+"),
    ]
    left, gap, width = 0.10, 0.08, 1.34
    for index, (day, head, body, seals) in enumerate(days):
        x = left + index * (width + gap)
        face = PALE_RED if index == 4 else PAPER
        edge = RED if index == 4 else BLUE
        patch = FancyBboxPatch((x, 0.25), width, 1.58,
                               boxstyle="round,pad=0.02,rounding_size=0.07",
                               fc=f"#{face}", ec=f"#{edge}", lw=1.1)
        ax.add_patch(patch)
        ax.text(x + 0.12, 1.65, day, ha="left", va="center", fontsize=7.5,
                fontfamily=FIG_SANS, fontweight="bold", color=f"#{edge}")
        ax.text(x + width / 2, 1.31, head, ha="center", va="center", fontsize=8.5,
                fontfamily=FIG_SANS, fontweight="bold", color=f"#{INK}")
        ax.text(x + width / 2, 0.91, body, ha="center", va="center", fontsize=7.4,
                color=f"#{MUTED}", linespacing=1.15)
        ax.text(x + width / 2, 0.48, f"target: {seals} seals", ha="center", va="center",
                fontsize=7.2, fontstyle="italic", color=f"#{edge}")
        if index < len(days) - 1:
            arrow(ax, (x + width, 1.03), (x + width + gap, 1.03), color=OCHRE, width=1.1)
    ax.plot([0.12, 7.02], [0.10, 0.10], color=f"#{BLUE}", lw=1.2)
    ax.text(0.12, 0.02, "quest forecast narrows", ha="left", va="top", fontsize=7,
            color=f"#{MUTED}")
    ax.text(7.02, 0.02, "VOTE  →  STORM  →  CONTINUE IN THE CHANGED CITY", ha="right", va="top",
            fontsize=6.8, fontfamily=FIG_SANS, fontweight="bold", color=f"#{RED}")
    return save_figure(fig, "quest_arc")


def quest_ownership_figure() -> Path:
    fig, ax = new_figure(7.2, 2.25)
    box(ax, 1.65, 1.20, 2.85, 1.55, "QUEST CODE VERIFIES",
        "clock • location • items • task receipts\nmeasure legality • conditions • seals • votes\nnotices • custody • quest outcome state",
        face=PALE_BLUE, edge=BLUE, title_color=BLUE, body_color=INK, title_size=9.4, body_size=7.25)
    box(ax, 5.55, 1.20, 2.85, 1.55, "SMART ACTORS INTERPRET",
        "whether they trust you • what they ask\nhow they read your words • whom they tell\nperformance, refusal, anger, mercy",
        face=PALE_OCHRE, edge=OCHRE, title_color=INK, body_color=INK, title_size=9.4, body_size=7.25)
    arrow(ax, (3.05, 1.45), (4.13, 1.45), color=RED, width=1.35)
    arrow(ax, (4.13, 0.95), (3.05, 0.95), color=RED, width=1.35)
    ax.text(3.60, 1.83, "quest-state seams", ha="center", va="center", fontsize=7.0,
            fontfamily=FIG_SANS, fontweight="bold", color=f"#{RED}")
    ax.text(3.60, 1.20, "propose condition\npledge / withdraw\nrecord testimony", ha="center", va="center",
            fontsize=7.0, color=f"#{RED}", linespacing=1.05,
            bbox=dict(fc=f"#{PAPER}", ec="none", pad=1.2))
    ax.text(3.60, 0.30, "The model may say no. It may never change the quest arithmetic.",
            ha="center", va="center", fontsize=8.2, fontstyle="italic", color=f"#{MUTED}")
    return save_figure(fig, "quest_ownership")


def quest_outcomes_figure() -> Path:
    fig, ax = new_figure(7.2, 2.35)
    outcomes = [
        ("ASSESSMENT", "low streets drain; every\nhousehold purse is lighter", BLUE, PALE_BLUE),
        ("CUT FEE", "repairs begin; cloth and\nfood carts queue and rise", OCHRE, PALE_OCHRE),
        ("POSTERN TOLL", "masonry holds; fish runs\nlate and smuggling grows", GREEN, "E1E8DE"),
        ("DEADLOCK", "cellars flood; wells close;\nemergency rule begins", RED, PALE_RED),
    ]
    for (head, body, color, face), x in zip(outcomes, [0.88, 2.70, 4.52, 6.34]):
        box(ax, x, 1.38, 1.55, 1.25, head, body, face=face, edge=color,
            title_color=color, body_color=INK, title_size=8.4, body_size=7.0)
    ax.text(3.60, 2.18, "THE NEXT MORNING IS THE QUEST EPILOGUE", ha="center", va="center",
            fontsize=9.5, fontfamily=FIG_SANS, fontweight="bold", color=f"#{INK}")
    ax.plot([0.18, 7.02], [0.56, 0.56], color=f"#{LINE}", lw=0.9)
    ax.text(0.24, 0.30, "HOW THIS QUEST WAS WON", ha="left", va="center", fontsize=7.2,
            fontfamily=FIG_SANS, fontweight="bold", color=f"#{MUTED}")
    ax.text(2.05, 0.30, "kept word", ha="center", va="center", fontsize=7.7, color=f"#{BLUE}")
    ax.text(3.42, 0.30, "compromised", ha="center", va="center", fontsize=7.7, color=f"#{OCHRE}")
    ax.text(4.89, 0.30, "forged / bribed", ha="center", va="center", fontsize=7.7, color=f"#{RED}")
    ax.text(6.42, 0.30, "who remembers it", ha="center", va="center", fontsize=7.7,
            fontstyle="italic", color=f"#{MUTED}")
    return save_figure(fig, "quest_outcomes")


def quest_milestones_figure() -> Path:
    fig, ax = new_figure(7.2, 1.55)
    steps = [
        ("M0", "ONE QUEST DAY", "deadline • roof\nquest card", BLUE),
        ("M1", "MINI-VOTE", "4 voters • 1 branch\nvisible rule change", RED),
        ("M2", "FULL QUEST", "16 voters • riders\n8 ward branches", OCHRE),
        ("M3", "AFTERMATH", "rumor • law routes\nnight receipts", GREEN),
        ("M4", "QUEST READY", "storm • persistence\naccessibility • polish", INK),
    ]
    for index, (code, head, body, color) in enumerate(steps):
        x = 0.74 + index * 1.43
        box(ax, x, 0.79, 1.16, 1.04, head, body, face=PAPER, edge=color,
            title_color=color, body_color=MUTED, title_size=7.5, body_size=6.55)
        ax.add_patch(Circle((x - 0.43, 1.30), 0.15, fc=f"#{color}", ec=f"#{color}"))
        ax.text(x - 0.43, 1.30, code, ha="center", va="center", fontsize=6.8,
                fontfamily=FIG_SANS, fontweight="bold", color="white")
        if index < len(steps) - 1:
            arrow(ax, (x + 0.58, 0.79), (x + 0.85, 0.79), color=LINE, width=1.0)
    return save_figure(fig, "quest_milestones")


def all_paragraphs(doc: Document):
    seen_cells: set[int] = set()

    def from_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    key = id(cell._tc)
                    if key in seen_cells:
                        continue
                    seen_cells.add(key)
                    yield from cell.paragraphs
                    yield from from_tables(cell.tables)

    yield from doc.paragraphs
    yield from from_tables(doc.tables)
    for section in doc.sections:
        for area in (section.header, section.footer):
            yield from area.paragraphs
            yield from from_tables(area.tables)


def replace_across_runs(paragraph, old: str, new: str) -> int:
    count = 0
    while old in paragraph.text:
        full = "".join(run.text for run in paragraph.runs)
        start = full.index(old)
        end = start + len(old)
        positions = []
        cursor = 0
        for index, run in enumerate(paragraph.runs):
            positions.append((index, cursor, cursor + len(run.text)))
            cursor += len(run.text)
        start_run = next(index for index, left, right in positions if left <= start < right)
        end_run = next(index for index, left, right in positions if left < end <= right)
        start_left = positions[start_run][1]
        end_left = positions[end_run][1]
        prefix = paragraph.runs[start_run].text[: start - start_left]
        suffix = paragraph.runs[end_run].text[end - end_left :]
        paragraph.runs[start_run].text = prefix + new + suffix
        for index in range(start_run + 1, end_run + 1):
            paragraph.runs[index].text = ""
        count += 1
    return count


REPLACEMENTS = [
    ("A shippable first campaign", "A shippable quest arc"),
    ("SIX-PAGE GDD", "SIX-PAGE QUEST OVERVIEW"),
    ("GAME DESIGN DOCUMENT", "QUEST OVERVIEW"),
    ("GAMEPLAY DESIGN", "QUEST DESIGN"),
    ("A first-person political immersive sim in Ombreval", "An authored five-day civic quest in Ombreval"),
    ("The sim verifies deeds, conditions and votes. Actors decide trust, terms and how refusal sounds.",
     "For this quest, code verifies deeds, conditions and votes; smart actors decide trust, terms and how refusal sounds."),
    ("Every favor means finding a real person, at a real office, before they move or sleep.",
     "Every quest favor means finding a real person, at a real office, before they move or sleep."),
    ("Speech reaches twenty metres. The useful sentence and the compromising one travel together.",
     "Quest-critical speech uses the existing twenty-metre hearing rule: useful and compromising words travel together."),
    ("Passing the measure is not a moral score. Its queues, prices, repairs and grudges become tomorrow’s play.",
     "Passing this measure is not a moral score. Its queues, prices, repairs and grudges become the quest’s aftermath."),
    ("FIRST CAMPAIGN", "FULL QUEST"),
    ("GAME DAYS", "QUEST DAYS"),
    ("HOW ONE DAY PLAYS", "HOW ONE QUEST DAY PLAYS"),
    ("How one day plays", "How one quest day plays"),
    ("Each day is a social heist against seven office bells.",
     "Each day of Nine Before Rain is a social heist against seven office bells."),
    ("The repeatable loop. Conversation discovers the opportunity; embodied play earns the receipt.",
     "The quest’s repeatable loop. Conversation discovers the opportunity; embodied play earns the receipt."),
    ("ONE DIEGETIC STATUS CARD", "ONE DIEGETIC QUEST CARD"),
    ("No persuasion percentage. No exhaustive quest arrows. The card records only public facts, explicit promises and leads the player has actually heard.",
     "No persuasion percentage. No exhaustive quest arrows. For this quest, the card records only public facts, explicit promises and leads the player has actually heard."),
    ("THE OFFICE LADDER", "EXISTING OFFICE LADDER — QUEST CLOCK (WATCH AT 02:00)"),
    ("The daily decision:", "The quest-day decision:"),
    ("FOUR JOB SHAPES; MANY HUMAN VERSIONS", "FOUR QUEST TASK SHAPES; MANY HUMAN VERSIONS"),
    ("FIVE DAYS, NINE SEALS", "FIVE QUEST DAYS, NINE PLEDGES"),
    ("Five days, nine seals", "Five quest days, nine pledges"),
    ("The campaign is one ugly arithmetic, not five hundred procedural quests.",
     "This is one authored coalition quest, not a procedural quest generator."),
    ("The exact storm is authored for the campaign. The forecast narrows, never ambushes the player with a hidden roll.",
     "The storm schedule is authored for this quest. The forecast narrows, never ambushes the player with a hidden roll."),
    ("Five days permit four or five ward knots, never all eight.",
     "The quest’s five-day deadline permits four or five ward branches, never all eight."),
    ("EIGHT AUTHORED WARD KNOTS", "EIGHT AUTHORED WARD BRANCHES FOR THIS QUEST"),
    ("WINNING THE NINTH", "SECURING NINE CONDITIONAL PLEDGES"),
    ("The microphone is a public act, not a menu-selection device.",
     "In this quest, the microphone is a public act, not a menu-selection device."),
    ("A playable day: visible service at the postern", "A playable quest day: visible service at the postern"),
    ("THE PLAYER VERBS", "QUEST ACTIONS AND SYSTEM DEPENDENCIES"),
    ("late-game sequence break", "late-quest optional route"),
    ("A critical reliability boundary: authored conditions and deterministic receipts preserve free conversation without letting a model invent or erase the rules.",
     "A critical reliability boundary for this quest: authored conditions and deterministic receipts preserve free conversation without letting a model invent or erase its rules."),
    ("Rohese Sedge will seal only a measure carrying a public tun audit. Deliver the true Gaudry weight by the Waning and she signs cleanly; substitute a borrowed weight and owe its owner; chalk a false proof and win the seal until a viewer catches it. The goal stays fixed. The route, witnesses, debt and aftermath are yours.",
     "Rohese Sedge will pledge only to a publicly accounted barrow toll with a neutral proving procedure. Bring the Line-keeper’s tally by the Waning and she agrees cleanly; borrow a witness and owe them; falsify the count and keep her only until a viewer catches it. The quest condition stays fixed. The route, witnesses, debt and aftermath are yours."),
    ("A seal is not favor points; it is a named promise with a visible predicate. Overnight memory or a new rumor may make a bencher withdraw, but the docket shows who, which condition broke, and what fact changed. A majority is therefore a route the player can reason about—not an invisible persuasion total.",
     "A seal tracks a conditional pledge, not favor points or the legal vote. If a predicate breaks, the docket names who withdrew and why. Only the public hand-count enacts the measure, so the coalition remains arithmetic the player can reason about—not an invisible persuasion total."),
    ("The goal stays fixed.", "The quest condition stays fixed."),
    ("There is no combat and no reload-shaped failure state.",
     "Nine Before Rain adds no combat, and none of its setbacks require reloading."),
    ("FIVE PRESSURES, ALL LEGIBLE", "FIVE QUEST PRESSURES, ALL LEGIBLE"),
    ("FAIL FORWARD", "QUEST FAIL-FORWARD ROUTES"),
    ("PROGRESSION WITHOUT XP", "LEVERAGE GAINED DURING THE QUEST"),
    ("The final upgrade is not a perk: it is a civic rule visible in tomorrow’s queues, prices, work routes and grudges.",
     "These are quest-state changes, not Cathedralbevy’s global progression model. The quest reward is an enacted civic rule visible in tomorrow’s queues, prices, work routes and grudges."),
    ("Passing the vote selects a civic consequence; how you passed it selects the social consequence.",
     "The passed motion selects the quest’s civic consequence; how you passed it selects the social consequence."),
    ("Nothing resets:", "No quest state resets:"),
    ("Persistence rule. Autosave at every office and before every explicit seal. After the hand-count, the same save continues into the storm morning; ‘ending’ means a changed civic state, never deletion of the city.",
     "Quest persistence requirement. Persist quest state at every office and before each explicit seal. After the hand-count, retain the outcome so play continues into the storm morning. ‘Quest complete’ means a changed civic state, never deletion of the city."),
    ("A SHIPPABLE FIRST CAMPAIGN", "A SHIPPABLE QUEST ARC"),
    ("Wrap the simulation in a deterministic civic spine; do not build a second game beside it.",
     "Add a bounded quest-state layer over the existing simulation."),
    ("FOUNDATION VS. NEW GAME LAYER", "EXISTING SYSTEMS VS. REQUIRED QUEST WORK"),
    ("SHIPS TODAY", "EXISTING SYSTEMS THIS QUEST USES"),
    ("NEW FOR NINE BEFORE RAIN", "FOUNDATION + QUEST WORK"),
    ("Five-day campaign state", "Five-day quest state"),
    ("Five-day quest state, persistence, safe roof and sleep-to-Kindling",
     "Quest state; save/load dependency; safe roof and sleep-to-Kindling"),
    ("Bounded Rumor Pollen; explicit morning memory / withdrawal receipts",
     "Political-rumor dependency; explicit morning withdrawal receipts"),
    ("60–90 MINUTE VERTICAL SLICE", "60–90 MINUTE QUEST SLICE"),
    ("first-time player", "first-time quest tester"),
    ("ending card", "quest-completion card"),
    ("an quest-completion card", "a quest-completion card"),
    ("CUT WITHOUT MERCY", "OUT OF SCOPE FOR THIS QUEST"),
    ("Design test:", "Quest acceptance test:"),
    ("Repository grounding.", "Quest dependencies."),
    ("Jonet will listen after you help divert the first dirty roof-water and inspect a cracked reserve seal.",
     "Jonet will listen after you operate Slate’s first-wash board and help make a witnessed tank-and-queue count."),
    ("You lift, carry and chalk the run-off while the queue watches. Code records place, time and witnesses.",
     "You work the board and tally buckets while the queue watches. Code records place, time and witnesses."),
    ("cash one more favor before the Waning, land one dangerous last word before a bencher sleeps, or turn back toward the clerk’s loft before the Scold makes your route illegal.",
     "cash one more favor before the Waning, land one dangerous last word before a bencher sleeps, or turn back before Snuffing closes ordinary meetings."),
    ("Shore the moving head and assign cost where both sides can watch.",
     "Witness shaft and windlass costs; secure a lawful hearing date."),
    ("Stage the Serle tun proving with chained Gaudry weights.",
     "Arrange a neutral proving procedure without settling the tun."),
    ("Forecast → first drops → clogged aprons → hard rain. The city visually rehearses what is at stake.",
     "Forecast → first drops → clogged aprons → authored downpour. Low streets and cellars rehearse what is at stake."),
    ("One curfew decision and one complete Stone House route through surety or escape.",
     "One night-route decision and one Stone House route through surety or escape."),
    ("Canonical Drain Question and wards: lore/core_lore/ward_politics.md; civic rules: secular_government.md; water work: lore/wells_and_water.md. Existing gameplay seams: cathedral-sim AGENTS.md and features/implemented/{law_and_order,chalking_the_walls,movement}. New bounded layers draw on features/{lore_ward_politics,rumors,false_peals__ring_the_bells_manually}.md.",
     "Canon: lore/core_lore/ward_politics.md, lore/core_lore/secular_government.md and lore/wells_and_water.md. Runtime: crates/cathedral-sim/AGENTS.md. Related specs: features/lore_ward_politics.md, features/knowledge_and_rumor/ and features/implemented/."),
    ("This quest does not add combat, a health tree, full procedural interiors, romance, free-form generated quests, magical powers, a 20,000-person crowd target, or solution to the Second Sun. False peals and the Cut game remain optional sequence breaks until the civic loop proves fun.",
     "This quest does not add combat, a health tree, full procedural interiors, romance, free-form generated quests, magical powers, a 20,000-person crowd target, or a solution to the Second Sun. Their presence or absence elsewhere in Cathedralbevy is outside this specification. False peals and the Cut game remain optional quest routes until the core civic loop proves fun."),
]


def replace_exact_cell_labels(paragraph):
    exact = {
        "SHIPS": "EXISTING SYSTEM",
        "SHIPS / CONTENT": "EXISTING SYSTEM + QUEST CONTENT",
        "NEW": "REQUIRED FOR QUEST",
        "NEW / OPTIONAL": "OPTIONAL QUEST EXTENSION",
    }
    current = paragraph.text.strip()
    if current not in exact:
        return
    first = next((run for run in paragraph.runs if run.text), None)
    for run in paragraph.runs:
        run.text = ""
    if first is not None:
        first.text = exact[current]


def reframe_text(doc: Document):
    paragraphs = list(all_paragraphs(doc))
    for paragraph in paragraphs:
        replace_exact_cell_labels(paragraph)
        for old, new in REPLACEMENTS:
            replace_across_runs(paragraph, old, new)

    # Replace the cover premise with an explicit quest-vs-game scope boundary.
    for paragraph in paragraphs:
        if paragraph.text.startswith("The pitch. Hard rain is coming") or paragraph.text.startswith("Quest premise. Hard rain is coming"):
            for run in paragraph.runs:
                run.text = ""
            label = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            label.text = "Document scope. "
            label.bold = True
            label.font.name = LABEL_FONT
            label._element.rPr.rFonts.set(qn("w:eastAsia"), LABEL_FONT)
            label.font.size = Pt(8.0)
            label.font.color.rgb = RGBColor.from_string(BLUE)
            body = paragraph.add_run(
                "Nine Before Rain is one optional five-day quest arc inside Cathedralbevy—not the whole game. "
                "Movement, conversation, schedules, law, inventory and persistence are dependencies. Hard rain "
                "is coming and the Common Bench has deadlocked. First Seat Gude Dask assigns an unknown visitor; "
                "the Common Clerk enters a narrow runner’s writ. Carry terms between eight hostile wards, do work "
                "people can witness, and return with nine conditional pledges for the legal hand-count. You cannot save "
                "everyone. You can make the cost legible, keep or break your word, and wake after the vote inside "
                "the city this quest changed."
            )
            body.font.name = BODY_FONT
            body._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
            body.font.size = Pt(9.2)
            body.font.color.rgb = RGBColor.from_string(INK)
            break

    # This paragraph was written as a whole-game exclusion. Scope it to this quest even if line wrapping
    # made the longer exact replacement above miss.
    for paragraph in paragraphs:
        if paragraph.text.startswith("No combat, health tree, full procedural interiors"):
            for run in paragraph.runs:
                run.text = ""
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = (
                "This quest does not add combat, a health tree, full procedural interiors, romance, free-form "
                "generated quests, magical powers, a 20,000-person crowd target, or a solution to the Second Sun. "
                "Their presence or absence elsewhere in Cathedralbevy is outside this specification. False peals "
                "and the Cut game remain optional quest routes until the core civic loop proves fun."
            )
            run.font.name = BODY_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
            run.font.size = Pt(8.35)
            run.font.color.rgb = RGBColor.from_string(INK)

    doc.core_properties.title = "Nine Before Rain — Quest Overview"
    doc.core_properties.subject = "Design overview for the Nine Before Rain quest in Cathedralbevy"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "quest design, civic quest, Ombreval, cathedralbevy"


def replace_inline_diagrams(doc: Document, paths: list[Path]):
    shapes = list(doc.inline_shapes)
    if len(shapes) != 7:
        raise RuntimeError(f"Expected 7 inline images in source overview, found {len(shapes)}")
    # 0 cover art, 1 loop, 2 arc, 3 triptych art, 4 ownership, 5 outcomes, 6 milestones.
    for shape_index, path in zip([1, 2, 4, 5, 6], paths):
        shape = shapes[shape_index]
        blip = shape._inline.graphic.graphicData.pic.blipFill.blip
        image_part = doc.part.related_parts[blip.embed]
        image_part._blob = path.read_bytes()


def build_docx():
    if not SOURCE.exists():
        raise FileNotFoundError(f"Generate the source six-pager first: {SOURCE}")
    diagrams = [
        quest_loop_figure(),
        quest_arc_figure(),
        quest_ownership_figure(),
        quest_outcomes_figure(),
        quest_milestones_figure(),
    ]
    doc = Document(SOURCE)
    reframe_text(doc)
    replace_inline_diagrams(doc, diagrams)
    doc.save(OUT_DOCX)


def convert_pdf():
    profile = Path(tempfile.mkdtemp(prefix="quest_overview_soffice_", dir="/tmp"))
    result = subprocess.run(
        [
            "soffice",
            "--headless",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(HERE),
            str(OUT_DOCX),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    if not OUT_PDF.exists():
        raise RuntimeError(f"soffice did not produce {OUT_PDF}\n{result.stdout}\n{result.stderr}")


def main():
    build_docx()
    convert_pdf()
    print(OUT_DOCX)
    print(OUT_PDF)


if __name__ == "__main__":
    main()
