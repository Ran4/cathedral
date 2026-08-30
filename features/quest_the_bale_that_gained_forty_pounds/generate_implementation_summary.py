#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx"]
# ///
"""Build the 3-4 page implementation summary for *Forty Pounds Over*.

The full design lives in ``README.md``; this document is only the part that has
to be built, for someone deciding whether to schedule it. Writes
``forty_pounds_over_implementation.docx`` beside this script.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
TARGET = HERE / "forty_pounds_over_implementation.docx"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x6B, 0x4A, 0x1F)


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.02

    for name, size, colour, before in (
        ("Heading 1", 16, ACCENT, 12),
        ("Heading 2", 12, ACCENT, 10),
        ("Heading 3", 10, INK, 6),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = colour
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def para(document: Document, text: str, *, style: str | None = None, size: float | None = None,
         italic: bool = False, colour: RGBColor | None = None, space_after: float | None = None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    if size:
        run.font.size = Pt(size)
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def rich(document: Document, chunks: list[tuple[str, str]], *, style: str | None = None):
    """Paragraph from (text, mode) chunks; mode is '', 'b', 'i', or 'c' (code)."""
    paragraph = document.add_paragraph(style=style)
    for text, mode in chunks:
        run = paragraph.add_run(text)
        run.font.bold = "b" in mode
        run.font.italic = "i" in mode
        if "c" in mode:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
    return paragraph


def bullets(document: Document, items: list[str], *, style: str = "List Bullet") -> None:
    for item in items:
        paragraph = document.add_paragraph(style=style)
        # Bold everything before the first em dash, which is the item's label.
        if " — " in item:
            label, rest = item.split(" — ", 1)
            paragraph.add_run(label).font.bold = True
            paragraph.add_run(" — " + rest)
        else:
            paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(2)


def table(document: Document, headers: list[str], rows: list[list[str]],
          widths: list[float] | None = None) -> None:
    control = document.add_table(rows=1, cols=len(headers))
    control.style = "Table Grid"
    control.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Word only honours explicit cell widths with autofit off, and it wants the
    # width on the column *and* on every cell in it.
    control.autofit = False
    header_cells = control.rows[0].cells
    for cell, text in zip(header_cells, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    for row in rows:
        cells = control.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            # A leading '`' marks the whole cell as code.
            if text.startswith("`") and text.endswith("`"):
                run = paragraph.add_run(text.strip("`"))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
            else:
                run = paragraph.add_run(text)
                run.font.size = Pt(8.5)
    if widths:
        for column, width in zip(control.columns, widths):
            column.width = Inches(width)
        for row in control.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)


def code(document: Document, lines: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(lines)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def build() -> None:
    document = Document()
    style_document(document)

    # ---------------------------------------------------------------- title
    title = para(document, "Forty Pounds Over", style="Heading 1", space_after=0)
    title.runs[0].font.size = Pt(22)
    para(document,
         "Implementation summary — features/quest_the_bale_that_gained_forty_pounds/",
         size=10, italic=True, colour=MUTED, space_after=2)
    para(document,
         "Quest 02 of features/systemic_quest_suggestions.md, promoted to spec 2026-08-30. "
         "Full design in README.md; ten worked routes in PLAYTHROUGHS.md. This document is "
         "only the part that has to be built.",
         size=9.5, colour=MUTED, space_after=10)

    # ------------------------------------------------------------- premise
    para(document, "What the quest is, in one paragraph", style="Heading 2")
    document.add_paragraph(
        "A corded bale of broadcloth weighs forty pounds more at the Wool Gate than it did on the "
        "Tallage beam two hours earlier. Officers call it smuggling, turn Hugh Crake's road cart "
        "around before a crowd, and impound the load until a public opening at the next Dayspring. "
        "Hugh is innocent; a crime did occur. Renn Crake, out of his depth on a salt forfeit and "
        "holding a promise he could not keep, diverted the cart through the bonded warehouse, cut "
        "the cord, added a weaver's off-book bolt and imitated the weigher's seal. The player has "
        "about a game day and a half to decide what the truth is worth, and to whom."
    )

    para(document, "The one design constraint that governs the build", style="Heading 2")
    rich(document, [
        ("The discrepancy is arithmetic, not a flag.", "b"),
        (" The bale's weight is the sum of what is physically inside it; the manifest is a separate "
         "written claim; the seal records what the beam called at sealing. Nothing anywhere stores "
         "“this bale is suspicious”. Remove the extra bolt and the bale genuinely weighs what its "
         "paper says. Add something else and the player has made it worse. This is what stops the "
         "quest being a dialogue tree with a stage, and it is why the resolution function reads no "
         "quest flags. Every engineering decision below follows from it — in particular, ", ""),
        ("there must never be an is_tampered boolean.", "b"),
    ])

    # --------------------------------------------------------- what exists
    para(document, "What already exists (so the scope is smaller than it looks)", style="Heading 2")
    para(document,
         "Nothing in the fiction has to be invented. Every named person is a shipped, authored "
         "character whose own sheet already contains their role; every location is a registered "
         "place with a baked nav pin; the cart is already in the data.",
         size=9.5, space_after=4)
    table(document,
          ["Already shipped", "What it gives the quest"],
          [
              ["`rounds.json → road_parties[0]`",
               "The brede_wool_gate party: Hugh Crake as leader, two carters, the Wool Gate, "
               "broadcloth as declared return cargo, and manifests that are already first-class "
               "objects in round.rs."],
              ["`notices.rs`",
               "The full accusation ladder — Word → Summoned → Warranted — with accused, wronged "
               "and taken fields, raise_notice / settle_notice, and a summons that names a bell. "
               "The pressure on Hugh is entirely existing code."],
              ["`custody.rs`",
               "Seizure, the escort tether, the Stone House, the posted fee and the fifth door out. "
               "Every route through arrest already works."],
              ["`items.json`",
               "cloth with grade ∈ {kersey, broadcloth} at 14/40 sparks; key, wax, ledger, deed, "
               "letter, tally_stick, writing_kit. badge already carries an "
               "authenticity: [counterfeit] domain — the precedent for a forged seal."],
              ["18 character sheets",
               "The weigher is going blind and has his daughter read the beam aloud; the carter "
               "knows every weight by the beam's note and cannot read; the notary is lying awake "
               "about a seal he failed to look twice at; the broker has already pledged his "
               "mother's warehouse keys. No sheet needs editing."],
              ["12 registered places",
               "The Tallage, its toll-house, the bonded warehouse, the bonded weighing yard, the "
               "Tally Bridge, the Wool Gate, the Draper's Reach, the Needle, the Stone House. No "
               "nav rebake."],
          ],
          widths=[1.75, 5.2])

    # ---------------------------------------------------------- what to build
    para(document, "What has to be built", style="Heading 2")

    para(document, "A. Three general capabilities (not quest-gated)", style="Heading 3")
    bullets(document, [
        "Item weight — an optional weight_lb on items.json kinds, keyed exactly like price_sparks "
        "so metadata variants differ (broadcloth 40 lb, kersey 24, wool bundle 28, grain sack 56). "
        "ItemCatalog::weight_lb(&Item), multiplied by stack quantity. Absent means zero; only lots "
        "and beams ever ask.",
        "crates/cathedral-sim/src/lots.rs — sealed, weighed consignments: real contents as item ids, "
        "a separate declared manifest, the weight the beam called at sealing, a seal impression, a "
        "cord state, and an append-only custody chain whose unrecorded leg is the crime's shape.",
        "The public weighing — Ewart Tarn's spoken procedure as a sim event: the standard cited, the "
        "figure called, who read it aloud, and everyone inside the ordinary 20 m hearing radius when "
        "it was. Witnesses reuse the existing single authoritative recipient calculation.",
    ])
    rich(document, [
        ("These three are the reason the quest is worth building now rather than later: suggestions "
         "13 (the fish cart), 14 (the locked grain store) and 16 (the disputed standard weight) all "
         "silently assume them, and become content problems once they exist.", "i"),
    ])

    para(document, "A hard prerequisite that is not this quest's to build", style="Heading 3")
    rich(document, [
        ("features/knowledge_and_rumor/", "c"),
        (" must land first, through its own M4. It owns ", ""),
        ("Fact", "c"), ("/", ""), ("holds()", "c"),
        (", the what_you_know prompt block, rumour propagation and garbling, the player's receipts, "
         "the journal overlay, and ", ""),
        ("World::arm_actor", "c"),
        (". Decided 2026-08-30 to ship both halves as one feature rather than an interim fact-only "
         "layer. The quest then authors JSON facts and owns no knowledge code at all \u2014 an earlier "
         "draft of this spec invented a per-quest knowledge enum, which is exactly the thing three "
         "quests each inventing one would have produced.", ""),
    ])
    rich(document, [
        ("The schedule cost is real and should be weighed here:", "b"),
        (" that feature's own principal risk is a perceptibility tuning pass \u2014 the LLM must voice an "
         "injected line often enough for the player to feel news travelling, without every mouth in a "
         "ward parroting the same sentence \u2014 and all three quest specs are now behind it.", ""),
    ])

    code(document,
         "impl Lot {\n"
         "    /// Sum of the contents' catalog weight. Never stored, always computed.\n"
         "    pub fn true_weight_lb(&self, world: &World) -> u32 { … }\n\n"
         "    /// What the opening will find. Positive means heavier than sealed.\n"
         "    pub fn discrepancy_lb(&self, world: &World) -> i64 {\n"
         "        self.true_weight_lb(world) as i64 - self.sealed_weight_lb as i64\n"
         "    }\n"
         "}")

    para(document, "B. Quest-specific", style="Heading 3")
    bullets(document, [
        "BaleQuest state on World behind an absent/default-off data gate: phase machine "
        "(Dormant → Arming → Seized → Investigating → OpeningDue → Opened → Resolved), the clock, "
        "per-character standing, dispositions, outcome. Receipts and the player's casebook are NOT "
        "here \u2014 they belong to the knowledge layer, because all three quests need them.",
        "The Arming phase, one game day before the seizure: the bolt minted as a real item in the "
        "weaver's workshop, the promise seeded as an authored fact, Renn's goal and round set through "
        "World::arm_actor, Ede's errand actually walked. Arming is played, not narrated \u2014 a player "
        "in the Wickmarket that day can watch it happen, so nobody needs to remember it.",
        "The quest-day leg override on brede_wool_gate — the shipped legs put the cart at the gate "
        "at Lamplight, which would leave the player only the night; on a quest day Hugh works ahead "
        "of his round to catch the Brede road, so the stop is at High Wick. Off when the quest is off.",
        "The public opening at the bonded warehouse: a pure function of lot state at that instant, "
        "reading no quest flags, with five branches on (contents vs declared, weight vs seal, seal "
        "authenticity, cord state).",
        "Five outcome packages applied to the live world — what the Draper's Reach pays by the ell, "
        "who brokers freight at the Tally Bridge, whether the Brede cart runs, queue length at the "
        "Tallage, and whether a real smuggling route with a copied key now exists.",
        "One new data file, assets/world/quests/bale.json \u2014 the lot, the seal, the cord, the "
        "custody chain, the outcome packages, and the quest's authored facts. Absent means the quest "
        "is off and golden prompts are byte-identical to today.",
    ])

    para(document, "C. Seams", style="Heading 3")
    rich(document, [
        ("Three new player commands, all about acting on the bale: ", ""),
        ("PlayerExamine", "c"),
        (" (ask a competent nearby character to read a seal, a cord or a beam), ", ""),
        ("PlayerStageWeighing", "c"), (" and ", ""), ("PlayerLodgeAmendment", "c"),
        (". One new message, EngineMessage::BaleQuest, on a dedicated monotonic quest revision.", ""),
    ])
    rich(document, [
        ("There is no AcceptQuest and no offer scene.", "b"),
        (" The two sibling quests are handed to the player by a named NPC with a writ; this one must "
         "not be, because Hugh Crake does not know he needs help and would not ask a stranger. The "
         "quest is walked into \u2014 by sight of the Wool Gate at High Wick, by earshot of a turned "
         "cart and frightened oxen, or by the word reaching the ward as ordinary garbled gossip. "
         "There is nothing to accept; what activates is the journal, on the first fact learned. The "
         "cost is that a player can miss it entirely, which is acceptable for a quest but would not "
         "be for a first one.", ""),
    ])
    rich(document, [
        ("Do not", "b"),
        (" put the journal in the actor/item PublicSnapshot or touch the public-state revision per "
         "fact learned — that republishes the whole cast and the configured crowd, and "
         "PublicSnapshot's 160 KiB bound already has little headroom. The Bevy side never re-derives "
         "weights, seal authenticity or guilt, and two fields are absolute: ", ""),
        ("Lot::seal.pressed_by", "c"),
        (" and ", ""), ("FactSource", "c"),
        (" — who pressed the wax, and why a fact is true — appear in no projection, prompt, log line, "
         "journal entry or HUD. A character who holds \u201cthe seal is an imitation\u201d knows the seal "
         "is wrong; they do not know whose hand it was.", ""),
    ])

    # ------------------------------------------------------------ milestones
    document.add_page_break()
    para(document, "Milestones", style="Heading 2")
    rich(document, [
        ("Prerequisite: features/knowledge_and_rumor/ through its own M4.", "b"),
        (" Nothing below starts before it.", ""),
    ])
    table(document,
          ["", "Milestone", "Contents", "Done when"],
          [
              ["M0", "Weight, lots, one clock",
               "weight_lb in the catalog; lots.rs; the shared quest host (phase, data gate, outcome "
               "application — not receipts); BaleQuest state and Arming; the leg override; the "
               "seizure; the opening and its resolution function — with no investigation content.",
               "Headless: start, advance to Dayspring, watch the bolt found and Hugh seized. "
               "Quest-disabled prompts, snapshots and goldens byte-identical."],
              ["M1", "The public weighing",
               "Weighing, standards, the reader seam, heard_by. Bertran Hobbe's proved weights as a "
               "real item and his Lowmarket proving as a real event. The whole lawful lane: "
               "reweigh, amend, pay duty, settle the notice.",
               "GO/NO-GO. If a lawful reweigh cannot be staged out of existing schedules and speech "
               "without a quest-only verb, the design is wrong and should be reworked here."],
              ["M2", "The evidence surface",
               "bale.json: the seven traces as authored facts with sealed holder sets and "
               "per-holder phrasings — no Rust. Arming wired. The three examinations (seal, cord, "
               "beam) as dispositions: who will say what they know. The scoped warehouse key. "
               "Deterministic fake-backend answers.",
               "Every trace reachable headless; FactSource and pressed_by appear in no rendered "
               "string."],
              ["M3", "The four lanes",
               "Social (three-in-one-room, Skell buying the bolt, the six households), covert "
               "(removing the bolt, re-laying the cord, forging a seal, altering the custody copy), "
               "predatory (selling the truth, taking the key). Notice escalation and custody "
               "integration, including the player being taken.",
               "Four materially different routes reach a cleared Hugh; correctness holds with the "
               "Night Office disabled."],
              ["M4", "Opening and aftermath",
               "The full public opening presentation; five outcome packages applied to the live "
               "world; follow-up lines that make the new state legible with no ending card.",
               "All five opening branches reachable from a scripted run; the aftermath is visible "
               "without opening a menu."],
              ["M5", "Content, polish, ship",
               "All six loom households, every failure row and recovery route, prompt tuning, a "
               "full golden re-bless, drive scripts per lane.",
               "Every failure-matrix row has a scripted run that continues to a resolution."],
          ],
          widths=[0.35, 1.15, 2.85, 2.6])

    # ---------------------------------------------------- dependencies/risks
    para(document, "Dependencies and decisions needed before M0", style="Heading 2")
    table(document,
          ["Item", "State", "Handling"],
          [
              ["Knowledge and rumour",
               "features/knowledge_and_rumor/ — SPEC ONLY",
               "HARD PREREQUISITE. Owns Fact/holds(), the what_you_know block, pollen propagation "
               "and garbling, player receipts, the journal and World::arm_actor. Decided 2026-08-30 "
               "to ship both halves as one feature. Its own M4 is an open-ended perceptibility "
               "tuning pass, and three quests are behind it."],
              ["Shared quest scaffolding",
               "Does not exist; three specs now each propose their own",
               "Land quest.rs in M0 as a phase, a data gate and outcome application. Smaller than it "
               "was: receipts and the casebook moved to the knowledge layer, so it no longer commits "
               "the sibling specs to much."],
              ["Keys and locked places",
               "features/keys_and_locked_places.md — SPEC ONLY",
               "M2 ships one authored key item opening exactly one door (the bonded warehouse) with "
               "a hard-coded reach check. Honest, but a second half-implementation of a spec'd "
               "feature — if the keys feature is close, do it first. The quest must not block on it."],
              ["settle_notice on a false accusation",
               "The ladder settles by returning `taken` or by the wronged party's say-so",
               "Neither fits “the accused did not do it”. M1 must decide whether amendment "
               "discharges a notice or whether a new discharge reason is needed. This is a change "
               "to shared law code, not quest code."],
              ["General credit path",
               "Only round::try_purchase moves sparks",
               "Duty payment and surety go through a narrow quest-owned settlement calling the same "
               "code. Third feature in a row to hit this (chalking, crowd knob). Worth fixing "
               "properly rather than working around a fourth time."],
              ["Memory / goal injection",
               "state.memories and state.goal are seed-only plus LLM-editable; no sim-side setter "
               "(actions.rs:2325, 2603)",
               "World::arm_actor lands with the knowledge layer. A seed, not an override — the "
               "actor's own set_goal/forget must win afterwards. Used for two characters here; if a "
               "fact can be a fact, it should be one."],
              ["Persistence",
               "No save exists",
               "The arc is deliberately a single session (~1.5 game days). Outcomes apply to the "
               "live world and are lost on quit, as everything else is."],
          ],
          widths=[1.45, 1.85, 3.65])

    para(document, "Two content risks", style="Heading 3")
    bullets(document, [
        "Two people are called Clemence — Clemence Hobbe the weaver (whose bolt it is) and Clemence "
        "Crake of the Tallage (Renn's mother, who owns the warehouse). Both are load-bearing and "
        "both are shipped. No rendered string may ever contain a bare “Clemence”; a test asserts it "
        "over every string the feature adds. An officer or an LLM confusing them mid-investigation "
        "is desirable; the author being confused is not.",
        "The smoothest route (persuading Ewart Skell to buy the bolt) quietly destroys the "
        "independence of the person the quest is about, and nothing tells the player. That is a "
        "deliberate design position, not an oversight — but it is the one thing in the feature that "
        "should be confirmed in playtest rather than assumed.",
    ])

    # ------------------------------------------------------------ acceptance
    para(document, "Acceptance criteria (abridged)", style="Heading 2")
    bullets(document, [
        "Determinism — cargo test -p cathedral-sim passes, all offline. With bale.json absent, "
        "golden prompts are byte-identical. No HashMap order reaches a snapshot, prompt or golden.",
        "The arithmetic property — a test removes one bolt from a sealed lot and asserts the "
        "discrepancy goes to zero with no other state touched. true_weight_lb is never cached.",
        "Coverage — all five opening branches reachable headless; at least four materially "
        "different routes clear Hugh, at least two of them exposing neither Clemence Hobbe nor "
        "Renn; every failure-matrix row continues to a resolution rather than a dead end.",
        "Cost — the stage cap and the single in-flight cognition slot are unchanged; the quest adds "
        "no LLM lane. PublicSnapshot size unchanged (the canary still passes). Outcome unaffected "
        "at CATHEDRAL_EXTRA_NPCS=2000.",
    ])

    para(document, "Vertical slice", style="Heading 3")
    code(document,
         "cargo run -p cathedral-backends --bin cathedral-headless -- \\\n"
         "    --fake --quest bale --start-office highwick \\\n"
         "    --seconds-per-day 600 --watch-clock 0.5")
    para(document,
         "Expected transcript: the gate stop and turnaround; a notice raised against rbrde; a "
         "weighing staged at the Tallage beam on Bertran Hobbe's proved weights before Odo Trask; "
         "an amendment lodged; forty sparks of duty paid; the notice settled; the Dayspring opening "
         "reduced to a formality; and three lines later, Ewart Skell at the Draper's Reach saying "
         "what he now pays by the ell.",
         size=9.5)

    document.save(TARGET)
    print(f"wrote {TARGET}")


if __name__ == "__main__":
    build()
