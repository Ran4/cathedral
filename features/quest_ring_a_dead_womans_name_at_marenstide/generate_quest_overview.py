# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx", "matplotlib", "pillow"]
# ///
"""Generate the GDD "One Bell Is Not a Name" (figures, .docx, .pdf).

    uv run docs/gdd/generate_gdd.py

Everything lands beside this script: figures/*.png, one_bell_is_not_a_name.docx and
.pdf (via `soffice --headless --convert-to pdf`).
"""
from __future__ import annotations

import os
import re
import subprocess
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
FIG_DIR = HERE / "quest_overview_figures"
FIG_DIR.mkdir(exist_ok=True)
REPO = HERE.parent.parent
MAP_SRC = REPO / "lore" / "places" / "ombreval_top_down_map_preview.png"
DOC_BASENAME = "one_bell_is_not_a_name_quest_overview"

scratch = os.environ.get(
    "GDD_SCRATCH",
    "/tmp/claude-1000/-home-ran-src-rust-cathedralbevy/c3f36f76-344e-48c0-9b76-ea2278698319/scratchpad",
)
os.environ.setdefault("MPLBACKEND", "Agg")
os.environ.setdefault("MPLCONFIGDIR", os.path.join(scratch, "mpl"))
os.makedirs(os.environ["MPLCONFIGDIR"], exist_ok=True)

import matplotlib  # noqa: E402

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
from matplotlib import font_manager  # noqa: E402
from matplotlib.patches import Circle, FancyArrowPatch, FancyBboxPatch  # noqa: E402
from PIL import Image  # noqa: E402

# ----------------------------------------------------------------------------- palette
PARCH = "#f4ecd8"
INK = "#2b2118"
ACC = "#9b2d20"  # brick red
MUTED = "#7a6a58"
NEW = "#6b8fa3"  # light blue: anything "new"

FONT_CANDIDATES = ["Liberation Serif", "DejaVu Serif"]
_available = {f.name for f in font_manager.fontManager.ttflist}
FIG_FONT = next((f for f in FONT_CANDIDATES if f in _available), "DejaVu Serif")
plt.rcParams.update(
    {
        "font.family": FIG_FONT,
        "text.color": INK,
        "axes.edgecolor": INK,
        "figure.facecolor": PARCH,
        "axes.facecolor": PARCH,
        "savefig.facecolor": PARCH,
    }
)
DPI = 200

# ================================================================================ TEXT
TITLE = "One Bell Is Not a Name"
DATE = "2026-08-27"

DOC = r"""
*Twenty-two years after the Hammering, the man who holds the drowned gives you his sister's name and no money. Six game days to the Marenstide reading: ninety-six sparks, one living person who will stand up and say she was who you say she was, and one line in a notary's book — and then thirty-four slow strokes, one a year of the life. Or the roll is read without her again, as it has been twenty-two times.*

## 1. What this is

**Document scope.** One optional six-day questline inside Cathedralbevy. Not the premise, not the main story, not a whole-game design. Movement, speech, items, the law ladder, the Night Office and persistence are dependencies it consumes and does not own. This overview was first written as a whole-game GDD; the reframing is itself the finding.

| 6 | 96 | 1 | 1 | 5 | 4–6 h |
|---|---|---|---|---|---|
| game days | sparks | witness standing | deed entered | outcomes | one playthrough |

{{fig:city_map|beside:1}}

**The premise is canon, played.** `lore/the_dry_boatmen.md` already establishes every load-bearing part of it: Saint Maren keeps the roll of the drowned; at Marenstide her roll is read at the church; since the diversion almost nobody drowns in Ombreval, so after F.415 the parish began adding the ward's Hammering dead to the roll of the drowned — *because the Hammering dead were denied their name-knells and this was a bell that could still be rung for them.* The traditional families object in the ward's flattest voice, **you cannot drown under slate**, and the wick-priest Renn Hobbe reads them anyway and takes the argument afterwards in the porch. The quest adds one name to that argument and a price to that bell.

**The cast is already carrying it.** None of this was written for a quest. Noll Fitch the sexton opens the ground, *rings Maren Smallvoice for the name-knell at one slow stroke a year of the life*, chalks the newest buried name on the charnel-door lintel, and remembers helping bury the Hammering dead at thirty-five — *the common toll was not a name-knell, whatever the Chapter wrote*. Renn Hobbe was twelve in the Hammering and watched his father climb toward the bells while stone fell: *I will not tell a family that one common toll settled their dead.* Cobb Hawser holds every name the Serle has taken in thirty years, in order, and *nobody has asked him for it in twenty years*. Tilman Rue stands at the back of paupers' burials watching who counts the strokes and who reads the lintel — *and who already knew it*. And some quiet purse has paid the paupers' ninety-six for years, which neither the sexton nor the priest has ever asked about. The quest is what happens when somebody finally asks.

**And the naming law supplies the difficulty.** `lore/families/family_rud.md` makes the Hawsers the worked example of a landed name — *a Hawser is nobody's concern, and the street lets the byname take him* — while the Ruds are the same law one stage on: a name with no property left, kept alive by nothing but **everyone else's paperwork**. So the quest is that law with a deadline on it. To get a Hawser into one book you must argue with a Rud who is already in six.

**What the player is promised.** Not a hero and not a detective: an errand-runner and a beggar of favours who must, before the feast, **raise the ninety-six** and put it in the sexton's hand, **find someone still living who remembers her** and get one of them to stand at the reading, **get her room into the notary's book** over a sitting tenant, and **be in earshot** when the bell rings.

**Three things it is made of.** *The price is fixed and the routes are not:* ninety-six is canon, and the four purses that can pay it mean four different things by the time the bell rings. *Code counts, people judge:* the sim owns the arithmetic, the kinship rule, the deadline and the receipts, and whether anyone helps a stranger is an LLM's judgement, never a rule. *Nothing is a marker:* a witness is found by saying a dead name within twenty metres of somebody who was alive in F.415, and caught before their next leg.

## 2. How it starts

Two scenes on the first day, both where the cast already stands.

**Cobb Hawser, at Tanners' Slip.** Sixty-three, a pauper, teasing junk cable back to tow from a stool, and — his sheet's own claim — **the man who holds the drowned**: every name the Serle has taken in thirty years, in order, with the weather and the boat and who was fool enough to send them out. He will tell the lot for a bowl at the Hungry Ox, and nobody has asked him in twenty years. He was born boat-family; the name went with the hull, and the street named him Hawser because rope is what he gathers now. He keeps every drowned name in his head and cannot get his own sister onto the parish's roll, because she died under a roof. He hands the letter over on the ordinary two-sided offer card — refusable — and says the thing the quest hangs on: *she was thirty-four, she died under the slate in F.415, and they never rang her.*

*Alis Hawser, my sister, thirty-four, of the room on the Cut. She died under the slate in F.415 and they never rang her. Ninety-six is what the parish asks for a bell and I have never had ninety-six. Noll Fitch will ring it. Renn Hobbe will read her at Marenstide with the drowned if somebody living will stand up and say she was there. I would, but I am her brother.*

**Noll Fitch, at the charnel door.** The sexton is fifty-seven, which makes him thirty-five in the Hammering year; he buried them. Everyone gets a name and a bell, he says, and that is the whole of his politics. He states the terms, and the terms are the quest: *the parish has one price for a bell, and it is the paupers' ninety-six, and I will take it from you as readily as from the purse I have never asked about. Renn Hobbe would read her tomorrow and cannot — the ward wants somebody living to stand up and say she was there, and he is the one who answers for it in the porch every year after. And the notary will not enter a room to a dead woman on one man's word — least of all her brother's.* Then he tells you how long the bell takes, in four words: **a stroke a year.** Thirty-four of them, slow, and Rue will count every one — and the bell that rings them is already in the code, comment and all.

**The first twenty minutes.** The crier's count starts at six. The casebook is the letter with what you have learned written on the back of it. Bitter Well hands out turns and a bucket carried to a widow court pays one; a carried message pays one; the fuller's yard on the Cut pays two, and the fuller's yard is where the room is. Within an office the player has said a dead woman's name to three people, watched one of them stop walking, and learned that the room has a man in it.

**Declining changes nothing and is a real option.** Marenstide arrives, Renn Hobbe reads the roll without her, and the letter can be offered again next year — which is precisely what has happened twenty-two times.

## 3. One quest day

{{fig:loop}}

**The office ladder is the quest clock.** Nothing here is new; the day already runs on seven bells, and the quest only says which of them matter.

| the Kindling 05 | Dayspring 07 | High Wick 12 | the Waning 15 | Lamplight 18 | the Snuffing 21 |
|---|---|---|---|---|---|
| yards and gates stir; the day's paid errands are given out here or not at all | the well queue: the cheapest place to be heard by five people at once | everyone is at a hearth, and the crier cries the count | the notary's last hour at the toll-house; deeds close | the last hour a Major hears you before sleeping on it | the Scold; gates shut, and the watch is the only audience left |

**The quest-day decision** is which of three things the Waning is spent on: earning the last four sparks, walking a householder to the toll-house, or catching the one witness who will be at her wharf until the bell. They are mutually exclusive most days, and that is the whole design.

**The microphone is a public act, not a menu.** You say a dead name to shake a memory loose; you say what a person told you to say, faithfully or with one word moved, and are paid on whether it was believed; you say something about a Rud within twenty metres of a grey coat and it is a summons. At the reading you say *Alis Hawser* where people are standing, and the sim does nothing with the sound at all — it only counts who was inside twenty metres to hear it.

## 4. Six days, three things

{{fig:arc}}

**The purse: ninety-six sparks.** Against the shipped catalogue that is eight coats or forty-eight loaves, in a city where ale, a herring and a chalk pen cost one each — deliberately more than honest errands produce in six days. The shortfall is the design, and the four ways to close it are not equivalent.

| Purse | How | What it costs |
|---|---|---|
| **Your own getting** | Authored errands paid on receipt: buckets, carried words, a fuller's load, a barrow to the Moorings. 1–4 each; a hard day is 8–12. | Days — and every hour earning is an hour not asking. |
| **A lender's book** | Averil Skell, the fish smoker of Maren's Green, lends at a spark on the dozen. It is a sideline on her sheet, not a trade. | A debt with a date. Averil has been known to sell a debt. |
| **The Custody's dry money** | Lise Copp's counter, where the grey clerks pay their quiet people. Tilman Rue already counts faces at paupers' burials and would like this one's. | Your name in a book you cannot read, and the funeral's faces in the Custody's. |
| **The quiet purse** | Somebody has paid the paupers' ninety-six for years and nobody has ever asked whose it is — and somebody meets in the sexton's crypt after a pauper's burial while he keeps the hinge oiled and himself elsewhere. Ask, and it pays. | On a condition: the name on the lintel becomes a pass, a word said at a door by people who need one. She stops being a woman and becomes a password. |

**The witness: six memories, one who will stand.** Six existing sheets are seeded with a memory of Alis, and every one of them was an adult or nearly so in F.415 — a content rule, not an accident.

| Who | Why they remember | Why they might not stand |
|---|---|---|
| **Nan**, washerwoman, 66 | Has washed for half the Cut so long she knows a household by its linen, and remembers every debt, birth and grudge on that street the way a rent-book does. | A pauper and a widow of the Weigh — and her daughter wears a grey coat at the chapter house, so standing up in a Reed argument is a family matter too. |
| **Gude**, herb-seller, 66 | Sixty-six years at the stall by the fish-hall steps; she has sold simples to that street her whole life. | One of the **Spared** — she looked up through the Great Rose and said plainly she saw one sun, and has been pitied and doubted since. She is tired of remembering for other people, and wants to know who is paying. |
| **Cobb Hawser**, rope-picker, 63 | Her brother, and the man who holds the drowned. | Kin, and the notary discounts kin. He knows it. It is why he asked you, and it is the tutorial for the rule. |
| **Hamel of the Reach**, bencher, 58 | Tallied freight into the dry warehouses on that cartway for thirty years; knows who held which door. | Born boat-family; the name went with the hull, and the ward seated him *because there are more Hawsers than Alders and the fallen like to see one of their own on a seat* — which is also why standing up for a fallen name in an election year is not free. |
| **Idonea Tarn**, bencher, 54 | Keeps the cooperage ledgers in a book that has never once been wrong; thirty-two when it happened. | Counts before she speaks, and wants to get through one more reckoning without being made to say an awkward word in public. |
| **Wyn Alder**, boatwoman, 34 | **Twelve when the Moorings roof came down in the Hammering and killed her father.** She took the fore-pole when her brother froze. | **The authored objection, and it is personal.** *Water keeps its name and its dead, and the sky can mind its own business.* Her own father died under a roof and she has never asked the parish to read him. If Alis is read, why not Alder? |

Wyn Alder's refusal on day three is the quest's best scene and its designed reversal, and none of it was written for the quest: her father, her line about water keeping its dead, and the name she speaks in the nave but not in the street are all on her sheet. She is not defending a tradition, she is defending her father from being made an exception to. The best way out is therefore not persuasion but **asking for Alder too** — get her to name her own dead to the priest and the objection inverts, the two names go up together, and the man whose father climbed toward the bells while stone fell is the last in the parish who will refuse that. Failing which: read the name *apart from* the roll in the porch argument Renn Hobbe takes every year; stand somebody else and let her say her piece publicly; or lean on her, since Averil Skell holds forty-eight sparks of her brother's debt and collects in the street — and the Night Office will have carried what you did to two wards by morning.

**The room: a deed twenty-two years late.** Alis held a room on the Cut. Tam Rud has it now — twenty-nine, a journeyman fuller, sworn to the Breach, seven years old when she died, and a **Rud**, which is to say a name that appears in the fulling-master's ell-tally, the ward book and a debt list, while a Hawser appears in nothing. The notary will enter a retrospective deed on two householders of two different wards, in person, before the Waning, neither of them kin, with no live counter-claim standing. He cannot be bought — that is on his sheet — so the routes are persuasion, settling Tam's debt, informing on him, or forging the ward's hand, which buys about two days because the ward overwrites any notice older than two game days with its own cross.

## 5. A day in the life

**Day four. Two days to Marenstide. Forty-one sparks.**

*The Kindling.* The fuller's yard on the Cut. Tam Rud gives you a load for two and asks, for the third morning, why a stranger keeps coming to the one room in the ward nobody wants. You do not tell him that the notary closes at the Waning.

*Dayspring.* Bitter Well. Sibbe Quern hands out turns, hand before trade, and gives you the Alders' bucket. Wyn Alder — who told you yesterday, flatly and in front of four people, that water keeps its name and its dead and the sky can mind its own business — takes it without looking at you and pays the spark anyway, because the bucket is the bucket. Her father died under a roof too. Neither of you says so.

*High Wick.* Jos Brant cries two days to Marenstide. Gude is at her stall by the fish-hall steps, sixty-six, and she remembers Alis perfectly well. What she wants to know before she says so in front of a notary is who is paying for the bell. It is a fair question, she has spent thirty years being asked to remember things for other people, and you do not have a clean answer yet.

*The Waning.* The Tallage. Odo Trask, over the open fee book: two householders of two different wards, in person, neither of them kin, before this bell. You have Nan of the Weigh. Hamel of the Reach is Reed, is not kin, tallied that cartway for thirty years — and is a sitting bencher in an election year, so his word has a price and he names it.

*Lamplight.* The Bell and Ladle. Renna Tapster banks the fire. Averil Skell will lend twelve at a spark on the dozen, due at the quarter, and mentions in the same breath that Tam Rud owes her eleven — which is a route to the room and a debt in the same sentence, and you cannot take both.

*The Snuffing.* Seven strokes, then the Scold. Tilman Rue is at the charnel door at somebody else's burial, doing what his sheet says he does — watching who counts the strokes and who reads the chalked name, and who already knew it. He counts you. Forty-one of ninety-six, one witness standing, no deed, two days.

*The Watch, unseen.* Gude settles being asked into a memory; the ward mood comes back *wary of the stranger who is buying something*. You will hear that tomorrow, in other people's mouths, before you have said a word.

## 6. What pushes back, and how it ends

| Pressure | What the player feels |
|---|---|
| **The count** | Six, five, four, cried daily at High Wick. Nothing else in the game has ever counted down. |
| **The offices** | People move, yards close, the notary goes home. A missed bell changes the route, not the objective. |
| **The stranger rule** | Your word weighs nothing until you have been named and remembered; most of the city does not think about you until you address it. |
| **Your own words** | Hearing has no walls. What you said within twenty metres of the wrong coat is a notice by morning, and the Night Office settles the rest into memories and ward moods. |
| **The law** | Word, summons, the Stone House, a posted three-spark fee. Arrest costs offices, never the quest. |

{{fig:outcomes}}

**Nothing here is a reload.** Miss a person and you learn tomorrow's round. Lose a witness and the casebook names the reason. Get arrested and the reading proceeds without you — which is one of the authored ways to end up *rung, not heard*. Come up short on the sixth day and the sexton will not ring on credit, but he will say so to your face, and Cobb will be standing there.

**Whose money rang it is tracked apart from whether it rang.** The integrity ledger records provenance, promises kept and broken, forged marks later discovered, and whether a man went to the Stone House to clear a room. It changes reactions, law and what people say afterwards. It is never a score, and it never moves the arithmetic after the bell.

## 7. What it would cost to build

**Most of it ships.** The quest consumes, unchanged: the clock and its seven offices; the microphone and typed chat as speech heard within twenty metres; two-sided offers and the item catalogue; the unknown-people rule with `knows` and `remember`; `go_to` and `places_known`; notices, custody and the Stone House; chalk marks with forging and scrubbing; homes and 1,101 nav doors; rounds and hearth meals; the Night Office's thirty Major reflections and eight ward moods; and the soundscape's assembled bell patterns, Saint Maren's small bell among them at 300 m.

**What the quest adds** is small, and shaped like the code already there.

| Quest work | Size | LLM spend |
|---|---|---|
| **Quest state and the six-day clock.** Catalog, phases, the reading window, receipts, the integrity ledger; all of it in the pure sim, off when the catalog is absent. | M | zero |
| **`ring_knell` and the lintel.** `BellPattern::NameKnell { years }` already ships — *one slow stroke per year of the life*, 300 m, with a logged stroke count. New: the sexton-only action that fires it, the one-slot `LintelName` mark, and a percept that carries the name. | S | one turn, bought by your offer |
| **`stand_witness` / `withdraw_witness` / `enter_deed`.** Typed actions validated against quest state; the model chooses which authored condition it names and whether it believes you. | S–M | one turn each |
| **Authored paid errands.** Four kinds — carry, spoken errand, attend with, statement to — each a receipt, not a job system. | M | reuses turns already bought |
| **The casebook.** A quest-only projection on its own revision; explicitly *not* in the actor snapshot. | S | zero |
| **Seeds.** The letter, six memories on existing sheets, the room and its tenant claim, the crier's count. | S | prompt lines only |

{{fig:milestones|beside:2}}

**Each milestone is playable.** M0 is the clock and a deterministic *unrung*; M1 is one witness, one purse and one bell at a reduced scale, and is the go/no-go for the whole feature; M2 is the full errand; M3 is consequence; M4 is the reading and its five outcomes; M5 and M6 are the foundation and the polish.

**The slice that decides it** is three days, one witness and a reduced sum: find Nan on her round, say a dead name inside twenty metres, hear what would move her, earn it, hand it over, and have her still standing there the next morning after the Night Office has had it. If asking one person to remember a dead woman and stand up for her is not compelling at that scale, the other five should not be authored.

**Three foundation dependencies, named honestly.** *Persistence:* there is no world save today, and a six-day quest cannot be one process lifetime. *Earning:* there is no wage, workplace or payout path anywhere in the sim — the quest owns four authored errands and a general economy stays a base-game feature. *Player curfew:* the Snuffing moves NPCs but no watch-witnessed detection exists, so night play is atmosphere until it does. Two named actors, the wick-priest and the notary, are Minor today and decisive here; promote them before shipping rather than treating a Minor as a Major quest subject.

**What we do not know yet.** Whether the model reliably surfaces a seeded memory when a dead name is said within twenty metres of it, or whether the name needs a code-side percept behind it — this is what M1 is for, and everything else waits on it. Whether ninety-six is reachable without the errands quietly becoming a job system. Whether Wyn Alder's principled refusal reads as a scene or as a broken flag. Whether being overwritten on the lintel by the next pauper is the point or a disappointment. And whether the parish's price should be waivable by the priest's own mercy, which would be the warmest moment in the quest and would also dissolve its only hard number.

**How it is proved.** `cathedral-headless --fake` runs the whole six days in seconds and asserts the arithmetic: ninety-six rings, ninety-five does not, a withdrawn witness never counts, kin never satisfies the deed, and the reading fires exactly once even if one pump crosses three offices. A `CATHEDRAL_HEADLESS=1` drive script then walks the slice in the real renderer and screenshots the lintel, with no window on anybody's desktop.

**Why this one and not another.** It is small, it is dated, it ends, and it leaves marks the next quest can find: a name on a lintel, a line in a book, a debt, and eight wards with an opinion about a stranger. That is the shape worth repeating — not because this errand is special, but because a city that can hold one of these can hold a dozen.

**Out of scope for this quest:** combat, health, skill trees, procedural characters, romance, full interiors, a quest generator, quest markers, a reputation number, and every part of the second sun and the impossible light. The Hammering's cause is canonically unestablished and stays that way; this quest is about a bell, a book and a name.
"""

# places the text sends the player to, as the plan index lists them (north +x, east -z)
MAP_PINS = [
    ("Saint Maren's: the reading and the bell", (-140.5, -265.6)),
    ("the Cut: the room, and the fuller in it", (-213.5, -63)),
    ("the Tallage toll-house: the notary's book", (-176.5, 73.5)),
    ("Bitter Well: buckets, queue, five listeners", (70.7, -280.7)),
    ("the Stone House: if a route goes wrong", (44.5, -207.2)),
]


CAPTIONS = {
    "city_map": "Where the quest sends you. The walled plan of Ombreval, pins: " + "; ".join(f"{i} {n}" for i, (n, _) in enumerate(MAP_PINS, 1)) + ".",
    "loop": "One quest day. The same six steps every day; the three things in the middle are what the sixth day counts, and nothing on the diagram is a HUD element.",
    "arc": "The six days of the quest: the authored beat of each day, and the purse, witness and room the player is expected to be holding by the end of it.",
    "outcomes": "How the quest ends. Every outcome is a position the city then lives in, and whose money rang the bell is tracked apart from whether it rang.",
    "milestones": "Build order. Each milestone is playable on its own; M1 is the vertical slice the rest of the quest is judged by.",
}
FIG_WIDTH_CM = {"city_map": 6.0, "loop": 11.2, "arc": 13.6, "outcomes": 13.4, "milestones": 6.3}


# ============================================================================== FIGURES
def box(ax, x, y, w, h, text, fc=PARCH, ec=INK, lw=1.2, ls="-", fs=8, weight="normal",
        color=None, pad=0.02, rounding=0.08, wrap=None, ha="center"):
    """Rounded box centred at (x, y) in data coords with text inside."""
    p = FancyBboxPatch((x - w / 2, y - h / 2), w, h,
                       boxstyle=f"round,pad={pad},rounding_size={rounding}",
                       fc=fc, ec=ec, lw=lw, ls=ls, mutation_aspect=1)
    ax.add_patch(p)
    if text:
        ax.text(x, y, text, ha=ha, va="center", fontsize=fs, weight=weight,
                color=color or INK, wrap=False, linespacing=1.15)
    return p


def titled_box(ax, x, y, w, h, head, body, fs_head=8.5, fs_body=7, **kw):
    box(ax, x, y, w, h, "", **kw)
    ax.text(x, y + 0.16, head, ha="center", va="center", fontsize=fs_head, weight="bold", color=INK)
    ax.text(x, y - 0.1, body, ha="center", va="center", fontsize=fs_body, color=INK, linespacing=1.15)


def arrow(ax, p0, p1, color=INK, lw=1.4, ls="-", label=None, lpos=0.5, loff=(0, 0),
          fs=7, style="-|>", shrinkA=6, shrinkB=6, connection="arc3,rad=0.0", lcolor=None,
          ha="center"):
    a = FancyArrowPatch(p0, p1, arrowstyle=style, color=color, lw=lw, ls=ls,
                        mutation_scale=11, shrinkA=shrinkA, shrinkB=shrinkB,
                        connectionstyle=connection)
    ax.add_patch(a)
    if label:
        lx = p0[0] + (p1[0] - p0[0]) * lpos + loff[0]
        ly = p0[1] + (p1[1] - p0[1]) * lpos + loff[1]
        ax.text(lx, ly, label, fontsize=fs, ha=ha, va="center", color=lcolor or color,
                style="italic", bbox=dict(fc=PARCH, ec="none", pad=1.0))
    return a


def new_axes(w_in, h_in):
    fig = plt.figure(figsize=(w_in, h_in))
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, w_in)
    ax.set_ylim(0, h_in)
    ax.set_aspect("equal")
    ax.axis("off")
    return fig, ax


def save(fig, name):
    out = FIG_DIR / f"{name}.png"
    fig.savefig(out, dpi=DPI)
    plt.close(fig)
    return out


def fig_loop():
    import math
    W, H = 6.7, 4.3
    fig, ax = new_axes(W, H)
    C = (3.35, 2.12)
    ax_r, ay_r = 2.38, 1.60
    nodes = [
        ("READ", "the crier's count, the casebook,\nwhat last night changed", 90),
        ("CHOOSE", "one or two leads; which purse,\nwhich name, which risk", 30),
        ("FIND", "their round, not a marker;\ncatch them before the next leg", -30),
        ("DO", "carry, work, pay, ask, witness,\nmark, inform, forge", -90),
        ("RETURN", "before their bell; a receipt,\na refusal, or a new lead", -150),
        ("NIGHT", "a roof, or the watch; the Night\nOffice settles what you said", 150),
    ]
    bw, bh = 1.9, 0.66
    pos = {}
    for name, body, deg in nodes:
        rad = math.radians(deg)
        x, y = C[0] + math.cos(rad) * ax_r, C[1] + math.sin(rad) * ay_r
        pos[name] = (x, y)
        titled_box(ax, x, y, bw, bh, name, body, fs_head=9.5, fs_body=6.7, lw=1.4)
    order = [n[0] for n in nodes]
    for i, name in enumerate(order):
        (x0, y0), (x1, y1) = pos[name], pos[order[(i + 1) % len(order)]]
        dx, dy = x1 - x0, y1 - y0
        d = math.hypot(dx, dy)
        sx, sy = dx / d, dy / d
        arrow(ax, (x0 + sx * bw / 2.3, y0 + sy * bh / 1.5), (x1 - sx * bw / 2.3, y1 - sy * bh / 1.5),
              lw=1.8, shrinkA=2, shrinkB=2, connection="arc3,rad=-0.14")
    ax.add_patch(Circle(C, 0.8, fc="#e8dcc6", ec=ACC, lw=1.7))
    ax.text(C[0], C[1] + 0.44, "BEFORE THE FEAST", ha="center", va="center", fontsize=7, color=ACC,
            weight="bold")
    ax.text(C[0], C[1] - 0.06, "96 sparks\na witness who will stand\nthe room in the book",
            ha="center", va="center", fontsize=7.6, color=ACC, linespacing=1.4)
    ax.text(C[0], C[1] - 0.55, "then one act", ha="center", va="center", fontsize=6.6, color=MUTED,
            style="italic")
    ax.text(0.05, H - 0.05, "one quest day", fontsize=7.5, color=MUTED, style="italic", rotation=90,
            ha="left", va="top")
    return save(fig, "loop")


def fig_arc():
    W, H = 6.7, 2.95
    fig, ax = new_axes(W, H)
    days = [
        ("DAY 1", "THE LETTER", "the charnel door;\nthree things named", "purse 0 · witness 0 / 6\nroom: unheard of"),
        ("DAY 2", "THE PRICE", "the room has a man\nin it; a fast purse", "purse 4-12 · witness 1\nroom: named"),
        ("DAY 3", "THE WAVERING", "a memory withdraws;\nthe notary wants two", "purse 12-28 · witness 1-2\nroom: one signature"),
        ("DAY 4", "THE SHORTFALL", "honest work cannot\nclose it in time", "purse 30-55 · witness 2\nroom: contested"),
        ("DAY 5", "THE TURN", "your words return\nfrom the Night Office", "purse 55-90 · witness 2\nroom: entered or lost"),
        ("DAY 6", "THE FEAST", "the knell rings;\nthe name said, or not", "96 · a witness stood\nthe room in the book"),
    ]
    x0, x1 = 0.12, W - 0.12
    n = len(days)
    gap = 0.1
    bw = (x1 - x0 - gap * (n - 1)) / n
    top, bh = H - 0.14, 1.92
    for i, (tag, head, body, target) in enumerate(days):
        x = x0 + i * (bw + gap)
        cx = x + bw / 2
        last = i == n - 1
        c = ACC if last else INK
        ax.add_patch(FancyBboxPatch((x, top - bh), bw, bh,
                                    boxstyle="round,pad=0,rounding_size=0.05",
                                    fc="#efe4cc" if last else PARCH, ec=c, lw=1.5 if last else 1.0))
        ax.text(cx, top - 0.18, tag, ha="center", va="center", fontsize=7, color=MUTED, weight="bold")
        ax.text(cx, top - 0.46, head, ha="center", va="center", fontsize=7.6, color=c, weight="bold")
        ax.text(cx, top - 0.92, body, ha="center", va="center", fontsize=6.6, color=INK, linespacing=1.2)
        ax.plot([x + 0.12, x + bw - 0.12], [top - bh + 0.6] * 2, color=MUTED, lw=0.5)
        ax.text(cx, top - bh + 0.32, target, ha="center", va="center", fontsize=6.1, color=MUTED,
                style="italic", linespacing=1.25)
        if not last:
            arrow(ax, (x + bw + 0.005, top - bh / 2), (x + bw + gap - 0.005, top - bh / 2),
                  lw=1.0, shrinkA=0, shrinkB=0)
    ax.text(W / 2, 0.68, "the deadline is authored; the route through it is not",
            ha="center", va="center", fontsize=7.5, style="italic", color=ACC)
    ax.plot([x0, x1], [0.44, 0.44], color=MUTED, lw=0.7)
    ax.text(x0, 0.26, "the letter is offered", ha="left", va="center", fontsize=7, color=MUTED, style="italic")
    ax.text(W / 2, 0.26, "one or two leads a day, caught on their own rounds", ha="center", va="center",
            fontsize=7, color=MUTED, style="italic")
    ax.text(x1, 0.26, "the city after", ha="right", va="center", fontsize=7, color=MUTED, style="italic")
    return save(fig, "arc")


def fig_outcomes():
    W, H = 6.7, 4.0
    fig, ax = new_axes(W, H)
    S = (0.55, 2.0)
    ax.add_patch(Circle(S, 0.44, fc=PARCH, ec=INK, lw=1.5))
    ax.text(S[0], S[1], "the letter\naccepted", ha="center", va="center", fontsize=7.5)
    mx, mw, mh = 2.5, 1.95, 0.56
    T, M, B = (mx, 2.9), (mx, 2.0), (mx, 1.1)
    box(ax, *T, mw, mh, "96 sparks in the\nsexton's hand", fs=8)
    box(ax, *M, mw, mh, "a witness who will\nstand and say it", fs=8)
    box(ax, *B, mw, mh, "the room entered\nin the notary's book", fs=8)
    for tgt in (T, M, B):
        arrow(ax, (S[0] + 0.44, S[1]), (tgt[0] - mw / 2, tgt[1]), shrinkA=1, shrinkB=2)
    purses = [("your own\ngetting", INK, "-"), ("a lender's\nbook", INK, "--"),
              ("the Custody's\ndry money", INK, "--"), ("the cell's\npurse", NEW, "--")]
    sx = [0.95, 2.0, 3.1, 4.2]
    lands = [mx - 0.72, mx - 0.24, mx + 0.24, mx + 0.72]
    for (t, c, ls), x, lx in zip(purses, sx, lands):
        box(ax, x, 3.72, 0.95, 0.42, t, fs=7, ec=c, color=c, ls=ls, lw=0.9)
        arrow(ax, (x, 3.51), (lx, T[1] + mh / 2), color=c, ls=ls, lw=0.8, shrinkA=1, shrinkB=1)
    ax.text(5.9, 3.72, "whose money rings it\ndecides what the name means",
            fontsize=6.8, style="italic", color=MUTED, ha="center", va="center", linespacing=1.2)
    ex, ew, eh = 5.85, 1.62, 0.5
    ys = [3.05, 2.42, 1.79, 1.16, 0.53]
    ends = [
        ("WELL RUNG: said in earshot,\ncarried out of the Night Office", dict(fc="#dfe6ea", ec=NEW, weight="bold", color=NEW)),
        ("RUNG, ROOM LOST: a name\nand nowhere it belongs", dict()),
        ("RUNG BY THE WRONG PURSE:\nthe name is somebody's asset", dict(weight="bold")),
        ("RUNG TO AN EMPTY YARD:\nchalked, unheard, overwritten", dict()),
        ("UNRUNG: the letter is\nstill in your coat", dict()),
    ]
    for (t, kw), y in zip(ends, ys):
        box(ax, ex, y, ew, eh, t, fs=7.2, lw=1.1, **kw)
    R, L = mx + mw / 2, ex - ew / 2
    for nd in (T, M, B):
        arrow(ax, (R, nd[1] + 0.12), (L, ys[0]), color=NEW, lw=0.9, shrinkA=1, shrinkB=1)
    arrow(ax, (R, T[1] - 0.14), (L, ys[1]), lw=0.9, shrinkA=1, shrinkB=1)
    arrow(ax, (R, M[1] - 0.16), (L, ys[3]), lw=0.9, shrinkA=1, shrinkB=1)
    px = sx[-1]
    arrow(ax, (px + 0.38, 3.72 - 0.21), (L, ys[2]), color=NEW, ls="--", lw=0.9, shrinkA=0, shrinkB=1,
          connection="angle,angleA=-90,angleB=180,rad=0")
    arrow(ax, (S[0], S[1] - 0.44), (L, ys[4]), lw=0.9, shrinkA=1, shrinkB=1,
          connection="angle,angleA=-90,angleB=180,rad=0")
    ax.text(1.95, 0.44, "the sixth day passes", fontsize=7.5, style="italic", ha="center", va="top")
    ax.text(W / 2, 0.1, "one save, no runs: every outcome is a civic state the city goes on living in",
            fontsize=7.5, style="italic", ha="center", va="center", color=MUTED)
    return save(fig, "outcomes")


def fig_milestones():
    W, H = 3.15, 4.82
    fig, ax = new_axes(W, H)
    items = [
        ("M0 State and the clock", "quest catalog, six-day deadline,\noffer / accept / decline, a\ndeterministic unrung ending",
         "playable:\naccept and\nrun out", INK, "-"),
        ("M1 One witness", "one memory, one paid task,\nring_knell + LintelName,\na mini-knell at three-day scale",
         "go / no-go\nslice", ACC, "-"),
        ("M2 The full errand", "six memories, the room, the\nregister, the casebook, all\nthree purses",
         "playable:\nthe whole\nerrand", INK, "-"),
        ("M3 Consequence", "night-office receipts, law and\ncustody routes, the occupant\nreacts, integrity ledger",
         "playable:\nyour own\nwords", INK, "-"),
        ("M4 The feast", "the office window, the name\nsaid, the sky, four aftermath\npackages",
         "playable:\nthe ending", NEW, "-"),
        ("M5 Persistence", "versioned checkpoint so a\nsix-day quest survives a\nsession (foundation)",
         "dependency,\nnot content", NEW, "--"),
        ("M6 Content and ship", "all routes authored, accessibility\nparity, audio and visual receipts,\nbalance pass",
         "shippable", NEW, "--"),
    ]
    bw = 2.28
    x = 1.19
    y = H - 0.05
    for i, (head, body, tag, c, ls) in enumerate(items):
        nl = body.count("\n") + 1
        bh = 0.24 + 0.118 * nl
        y -= bh / 2
        box(ax, x, y, bw, bh, "", ec=c, ls=ls, lw=1.3)
        ax.text(x - bw / 2 + 0.06, y + bh / 2 - 0.11, head, fontsize=8.2, weight="bold", ha="left",
                va="center", color=c)
        ax.text(x - bw / 2 + 0.06, y - 0.09, body, fontsize=6.6, ha="left", va="center", color=INK,
                linespacing=1.1)
        ax.text(x + bw / 2 + 0.06, y, tag, fontsize=6.6, ha="left", va="center", style="italic",
                color=MUTED, linespacing=1.05)
        y -= bh / 2
        if i < len(items) - 1:
            arrow(ax, (x, y), (x, y - 0.09), lw=1.2, shrinkA=0, shrinkB=0)
            y -= 0.09
    return save(fig, "milestones")


def fig_city_map():
    from PIL import ImageDraw, ImageFont

    out = FIG_DIR / "city_map.png"
    im = Image.open(MAP_SRC).convert("RGB")
    # fitted against the index: the Lanthorn (0, -12) and the Gradine (0, 131) on the preview
    col0, row0, k = 1264, 1466, 2.67

    def px(x, z):
        return (col0 - z * k, row0 + x * k)

    crop = (300, 460, 2560, 2440)  # the walled plan; the key and index live in the right third
    im = im.crop(crop)
    draw = ImageDraw.Draw(im)
    try:
        font = ImageFont.truetype(font_manager.findfont(FIG_FONT), 54)
    except Exception:
        font = ImageFont.load_default()
    r = 40
    for i, (_, (x, z)) in enumerate(MAP_PINS, 1):
        cx, cy = px(x, z)
        cx, cy = cx - crop[0], cy - crop[1]
        draw.ellipse((cx - r, cy - r, cx + r, cy + r), fill=ACC, outline=PARCH, width=6)
        draw.text((cx, cy), str(i), fill=PARCH, font=font, anchor="mm")
    w = 1400
    h = round(im.height * w / im.width)
    im.resize((w, h), Image.LANCZOS).save(out, optimize=True)
    return out


def draw_all():
    return {
        "loop": fig_loop(),
        "arc": fig_arc(),
        "outcomes": fig_outcomes(),
        "milestones": fig_milestones(),
        "city_map": fig_city_map(),
    }


# ================================================================================ DOCX
from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402,F401
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402

BODY_FONT = "Liberation Serif"
MONO_FONT = "Liberation Mono"
INK_RGB = RGBColor(0x2B, 0x21, 0x18)
ACC_RGB = RGBColor(0x9B, 0x2D, 0x20)
MUTED_RGB = RGBColor(0x7A, 0x6A, 0x58)
BODY_PT = 10.5
TABLE_PT = 7


def set_font(run_or_style, name=BODY_FONT, size=None, bold=None, italic=None, color=None):
    f = run_or_style.font
    f.name = name
    rpr = run_or_style._element.get_or_add_rPr() if hasattr(run_or_style, "_element") else run_or_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = color


def para_fmt(p, before=0, after=2, line=1.11, keep_next=False, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    if align is not None:
        p.alignment = align


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def add_inline(p, text, size=BODY_PT, base_italic=False, base_bold=False, color=None):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        bold, italic, mono = base_bold, base_italic, False
        if part.startswith("**") and part.endswith("**"):
            bold, part = True, part[2:-2]
        elif part.startswith("*") and part.endswith("*") and len(part) > 1:
            italic, part = not base_italic, part[1:-1]
        elif part.startswith("`") and part.endswith("`"):
            mono, part = True, part[1:-1]
        # a code span nested inside a bold or italic run keeps the run's weight
        for sub in re.split(r"(`[^`]+`)", part) if not mono else [part]:
            if not sub:
                continue
            sub_mono = mono or (sub.startswith("`") and sub.endswith("`") and len(sub) > 1)
            if sub_mono and not mono:
                sub = sub[1:-1]
            r = p.add_run(sub)
            set_font(r, MONO_FONT if sub_mono else BODY_FONT, size=size - (0.8 if sub_mono else 0),
                     bold=bold, italic=italic, color=color)


def set_cell_shading(cell, hex_fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def set_table_borders(table, color="2B2118", sz=4, none=False):
    tbl = table._tbl
    tblpr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if none:
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    tblpr.append(borders)


def set_cell_margins(table, top=10, bottom=10, left=55, right=55):
    tblpr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblpr.append(mar)


def cant_split(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    trpr.append(el)


def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_para.add_run("Page ")
    set_font(r, size=8.5, color=MUTED_RGB)
    r = footer_para.add_run()
    set_font(r, size=8.5, color=MUTED_RGB)
    for tag, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if tag:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        r._r.append(el)


class Builder:
    def __init__(self, figures):
        self.figures = figures
        self.fig_n = 0
        self.doc = Document()
        d = self.doc
        sec = d.sections[0]
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(sec, side, Cm(2.0))
        sec.footer_distance = Cm(1.0)
        sec.header_distance = Cm(1.0)
        st = d.styles["Normal"]
        set_font(st, size=BODY_PT, color=INK_RGB)
        st.paragraph_format.space_after = Pt(2)
        st.paragraph_format.line_spacing = 1.15
        for name, size, before, after in (("Heading 1", 13, 5, 2), ("Heading 2", 11.5, 4, 2)):
            hs = d.styles[name]
            set_font(hs, size=size, bold=True, color=ACC_RGB if name == "Heading 1" else INK_RGB)
            hs.paragraph_format.space_before = Pt(before)
            hs.paragraph_format.space_after = Pt(after)
            hs.paragraph_format.keep_with_next = True
            hs.paragraph_format.line_spacing = 1.05
        for name in ("List Bullet", "List Number"):
            ls = d.styles[name]
            set_font(ls, size=BODY_PT, color=INK_RGB)
            ls.paragraph_format.space_after = Pt(1)
            ls.paragraph_format.line_spacing = 1.15
        add_page_number(sec.footer.paragraphs[0])

    # ---- blocks
    def title_block(self):
        d = self.doc
        p = d.add_paragraph()
        para_fmt(p, after=2, line=1.0)
        r = p.add_run(f"CATHEDRALBEVY   ·   QUEST OVERVIEW   ·   SIX PAGES   ·   {DATE}")
        set_font(r, size=8, bold=True, color=MUTED_RGB)
        p = d.add_paragraph()
        para_fmt(p, after=1, line=1.0)
        r = p.add_run(TITLE)
        set_font(r, size=24, bold=True, color=ACC_RGB)
        p = d.add_paragraph()
        para_fmt(p, after=6, line=1.0)
        r = p.add_run("A six-day errand in Ombreval")
        set_font(r, size=12, italic=True, color=INK_RGB)
        r = p.add_run("   ·   features/quest_ring_a_dead_womans_name_at_marenstide")
        set_font(r, size=8.5, color=MUTED_RGB)
        # rule
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "2B2118")
        pbdr.append(bottom)
        pPr.append(pbdr)

    def heading(self, text, level):
        p = self.doc.add_paragraph(style=f"Heading {level}")
        r = p.add_run(text)
        set_font(r, size=13.5 if level == 1 else 11.5, bold=True,
                 color=ACC_RGB if level == 1 else INK_RGB)
        p.paragraph_format.keep_with_next = True

    def paragraph(self, text, keep_next=False):
        p = self.doc.add_paragraph()
        para_fmt(p, keep_next=keep_next, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        add_inline(p, text)
        return p

    def list_item(self, text, style):
        p = self.doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(1)
        add_inline(p, text)

    def table(self, header, rows):
        d = self.doc
        ncols = len(header)
        t = d.add_table(rows=1, cols=ncols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        set_table_borders(t)
        set_cell_margins(t)
        total = 17.0
        # Column widths follow the content: a column of "S"/"M" must not be handed the
        # same centimetres as a column of sentences, and a first column of names must not
        # eat the width the explanation needs. Weight by the longest cell (which is what
        # actually has to wrap), softened by the mean so one outlier cannot dominate.
        cols = [[header[i]] + [r[i] for r in rows if i < len(r)] for i in range(ncols)]
        weights = []
        for c in cols:
            lens = [len(str(x)) for x in c] or [1]
            weights.append(0.65 * max(lens) + 0.35 * (sum(lens) / len(lens)))
        floor = 1.35 if ncols > 3 else 2.0
        widths = [total * w / sum(weights) for w in weights]
        for _ in range(4):  # lift short columns to the floor, take it from the rest pro rata
            short = [i for i, w in enumerate(widths) if w < floor]
            if not short:
                break
            debt = sum(floor - widths[i] for i in short)
            free = sum(widths[i] for i in range(ncols) if i not in short)
            if free <= debt:
                widths = [total / ncols] * ncols
                break
            for i in range(ncols):
                widths[i] = floor if i in short else widths[i] - debt * widths[i] / free
        for i, cell in enumerate(t.rows[0].cells):
            cell.width = Cm(widths[i])
            set_cell_shading(cell, "E8DCC6")
            p = cell.paragraphs[0]
            para_fmt(p, after=0, line=1.0, keep_next=True)
            add_inline(p, header[i], size=TABLE_PT, base_bold=True)
        cant_split(t.rows[0])
        # repeat the header row when the table breaks across a page (tblHeader; the
        # python-docx property of that name is not in every release)
        t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for row in rows:
            cells = t.add_row().cells
            for i, txt in enumerate(row):
                cells[i].width = Cm(widths[i])
                p = cells[i].paragraphs[0]
                para_fmt(p, after=0, line=1.0)
                add_inline(p, txt, size=TABLE_PT)
            cant_split(t.rows[-1])
        # column widths must be set on the grid, the columns and every cell for LibreOffice
        for i, col in enumerate(t.columns):
            col.width = Cm(widths[i])
        for row in t.rows:
            for i, c in enumerate(row.cells):
                c.width = Cm(widths[i])
        grid = t._tbl.find(qn("w:tblGrid"))
        for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
            gc.set(qn("w:w"), str(int(widths[i] / 2.54 * 1440)))
        sp = d.add_paragraph()
        para_fmt(sp, after=2, line=1.0)
        sp.paragraph_format.space_after = Pt(2)
        for r in sp.runs:
            set_font(r, size=4)
        sp.add_run().font.size = Pt(4)

    def figure_cell(self, container, fig_id, width_cm, caption_in=None):
        """Add picture + caption into a paragraph flow `container` (doc or cell).

        `caption_in` lets the caption land in a different container (the beside layout
        puts it under the whole table at full width so it stays one line)."""
        self.fig_n += 1
        p = container.add_paragraph() if container is self.doc else container.paragraphs[0]
        # keep_with_next inside a table cell makes LibreOffice keep the whole table with the
        # following paragraph, which can push the table to a new page; the caption follows
        # the table there anyway.
        para_fmt(p, before=1, after=0, line=1.0, keep_next=caption_in is None,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run().add_picture(str(self.figures[fig_id]), width=Cm(width_cm))
        if caption_in is None:
            self.caption(container, fig_id)

    def caption(self, container, fig_id):
        cap = container.add_paragraph()
        para_fmt(cap, after=3, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = cap.add_run(f"Figure {self.fig_n} — {CAPTIONS[fig_id]}")
        set_font(r, size=8.5, italic=True, color=MUTED_RGB)

    def figures_block(self, ids):
        if len(ids) == 1:
            fid = ids[0]
            w = FIG_WIDTH_CM[fid]
            if w >= 12:
                self.figure_cell(self.doc, fid, w)
            else:
                self.side_by_side([fid])
        else:
            self.side_by_side(ids)

    def beside(self, fid, paragraphs):
        """Half-width figure on the left, the given paragraphs flowing on the right."""
        t = self.doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        set_table_borders(t, none=True)
        set_cell_margins(t, top=0, bottom=0, left=0, right=80)
        widths = [FIG_WIDTH_CM[fid] + 0.2, 17.0 - FIG_WIDTH_CM[fid] - 0.2]
        left, right = t.rows[0].cells
        for i, c in enumerate((left, right)):
            c.width = Cm(widths[i])
            t.columns[i].width = Cm(widths[i])
        grid = t._tbl.find(qn("w:tblGrid"))
        for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
            gc.set(qn("w:w"), str(int(widths[i] / 2.54 * 1440)))
        self.figure_cell(left, fid, FIG_WIDTH_CM[fid], caption_in=self.doc)
        # not cant-split: the text column may run on to the next page
        first = True
        for text in paragraphs:
            p = right.paragraphs[0] if first else right.add_paragraph()
            first = False
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            add_inline(p, text)
        # the caption is a merged full-width second row, so it stays one line and travels
        # with the picture instead of being orphaned under a page break
        cap_row = t.add_row()
        merged = cap_row.cells[0].merge(cap_row.cells[1])
        cap = merged.paragraphs[0]
        para_fmt(cap, before=1, after=3, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = cap.add_run(f"Figure {self.fig_n} — {CAPTIONS[fid]}")
        set_font(r, size=8.5, italic=True, color=MUTED_RGB)
        cant_split(cap_row)
        # a table may still break *between* its rows, which orphans the caption at the top
        # of the next page. keepNext on every paragraph of the picture row forbids that.
        for c in t.rows[0].cells:
            for p in c.paragraphs:
                p.paragraph_format.keep_with_next = True

    def side_by_side(self, ids):
        t = self.doc.add_table(rows=1, cols=len(ids))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t, none=True)
        for cell, fid in zip(t.rows[0].cells, ids):
            cell.width = Cm(17.0 / len(ids))
            self.figure_cell(cell, fid, FIG_WIDTH_CM[fid])
        cant_split(t.rows[0])
        sp = self.doc.add_paragraph()
        para_fmt(sp, after=0, line=1.0)
        sp.add_run().font.size = Pt(2)

    # ---- markdown
    def render(self, md):
        lines = md.strip("\n").split("\n")
        i = 0
        para_buf = []

        def flush():
            nonlocal para_buf
            if para_buf:
                self.paragraph(" ".join(para_buf))
                para_buf = []

        while i < len(lines):
            line = lines[i]
            s = line.strip()
            if not s:
                flush()
                i += 1
                continue
            m = re.match(r"^(#{1,3})\s+(.*)$", s)
            if m:
                flush()
                self.heading(m.group(2), len(m.group(1)))
                i += 1
                continue
            mb = re.match(r"\{\{fig:(\w+)\|beside:(\d+)\}\}", s)
            if mb:
                flush()
                fid, count = mb.group(1), int(mb.group(2))
                paras = []
                i += 1

                def _flowable(ln):
                    """Only body paragraphs may be pulled in beside a figure; a heading,
                    table or figure that followed would otherwise be swallowed and printed
                    as literal markdown."""
                    return bool(ln) and not ln.startswith(("#", "|", "{{fig:", "- ", "* "))

                while len(paras) < count and i < len(lines):
                    ln = lines[i].strip()
                    if not ln:
                        i += 1
                        continue
                    if not _flowable(ln):
                        break
                    buf = []
                    while i < len(lines) and _flowable(lines[i].strip()):
                        buf.append(lines[i].strip())
                        i += 1
                    paras.append(" ".join(buf))
                self.beside(fid, paras)
                continue
            if s.startswith("{{fig:"):
                flush()
                ids = [re.match(r"\{\{fig:(\w+)\}\}", s).group(1)]
                # consecutive half-width figure placeholders pair up
                while i + 1 < len(lines) and lines[i + 1].strip().startswith("{{fig:") and \
                        FIG_WIDTH_CM[ids[-1]] < 12:
                    nid = re.match(r"\{\{fig:(\w+)\}\}", lines[i + 1].strip()).group(1)
                    if FIG_WIDTH_CM[nid] >= 12:
                        break
                    ids.append(nid)
                    i += 1
                self.figures_block(ids)
                i += 1
                continue
            if s.startswith("|"):
                flush()
                rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    if not all(re.fullmatch(r":?-+:?", c) for c in cells):
                        rows.append(cells)
                    i += 1
                self.table(rows[0], rows[1:])
                continue
            m = re.match(r"^(\d+)\.\s+(.*)$", s)
            if m:
                flush()
                self.list_item(m.group(2), "List Number")
                i += 1
                continue
            if s.startswith("- ") or s.startswith("* "):
                flush()
                self.list_item(s[2:], "List Bullet")
                i += 1
                continue
            para_buf.append(s)
            i += 1
        flush()

    def save(self, path):
        self.doc.save(path)


def build_docx(figures):
    b = Builder(figures)
    b.title_block()
    b.render(DOC)
    out = HERE / f"{DOC_BASENAME}.docx"
    b.save(out)
    return out


def to_pdf(docx_path):
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(HERE), str(docx_path)],
                   check=True, capture_output=True)
    return HERE / f"{DOC_BASENAME}.pdf"


def main():
    figs = draw_all()
    print("figures:", ", ".join(str(p) for p in figs.values()))
    docx_path = build_docx(figs)
    print("docx:", docx_path)
    if "--no-pdf" in sys.argv:
        return
    pdf = to_pdf(docx_path)
    print("pdf:", pdf)
    try:
        n = subprocess.run(["pdfinfo", str(pdf)], capture_output=True, text=True).stdout
        pages = re.search(r"Pages:\s+(\d+)", n)
        if pages:
            print("pages:", pages.group(1))
    except FileNotFoundError:
        pass


if __name__ == "__main__":
    main()
